import re
import io
import PyPDF2
from datetime import datetime, timedelta
from dataclasses import dataclass
from typing import List, Optional
from docx import Document
import os

@dataclass
class ChainEntry:
    date: datetime
    date_string: str
    grantor: str
    grantee: str
    instrument: str
    book_page: str
    remark: str = ""
    is_vesting: bool = False
    line: str = ""

def extract_text_from_pdf(pdf_stream) -> str:
    """Extract text from PDF stream using pdfplumber first, then fall back to PyPDF2."""
    text = ""
    # Try pdfplumber for more reliable layout text
    try:
        import pdfplumber
        with pdfplumber.open(pdf_stream) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}")

    if text.strip():
        return text

    # Fallback: PyPDF2
    try:
        pdf_stream.seek(0) # Reset stream for fallback
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"PyPDF2 extraction failed: {e}")
    return text

def extract_table_entries_from_pdf(pdf_stream) -> List[ChainEntry]:
    """Extract chain entries from PDF tables using a stream."""
    entries: List[ChainEntry] = []
    try:
        import pdfplumber
    except Exception:
        return entries

    table_settings = {
        "vertical_strategy": "lines",
        "horizontal_strategy": "lines",
        "intersection_tolerance": 5,
        "snap_tolerance": 3,
        "join_tolerance": 3,
        "edge_min_length": 3,
        "min_words_vertical": 1,
        "min_words_horizontal": 1,
    }

    try:
        with pdfplumber.open(pdf_stream) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables(table_settings=table_settings)
                if not tables:
                    continue
                for tbl in tables:
                    if not tbl or not any(tbl):
                        continue
                    # Find header row mapping
                    header_idx = None
                    col_map: dict[str, int] = {}
                    for r_idx, row in enumerate(tbl):
                        if not row:
                            continue
                        header_cells = [ (c or '').strip().upper() for c in row ]
                        if ('GRANTOR' in ' '.join(header_cells) and
                            'GRANTEE' in ' '.join(header_cells) and
                            'INSTRUMENT' in ' '.join(header_cells)):
                            header_idx = r_idx
                            # Build column index map by nearest match
                            def find_col(name_variants: List[str]) -> int:
                                for nv in name_variants:
                                    if nv in header_cells:
                                        return header_cells.index(nv)
                                # Try partial match
                                for nv in name_variants:
                                    for i, cell in enumerate(header_cells):
                                        if nv in cell:
                                            return i
                                return -1
                            col_map = {
                                'grantor': find_col(['GRANTOR']),
                                'grantee': find_col(['GRANTEE']),
                                'instrument': find_col(['INSTRUMENT']),
                                'date': find_col(['DATED', 'FILED', 'DATE']),
                                'recording': find_col(['BOOK-PAGE', 'RECORDING', 'BOOK', 'RECORD']),
                            }
                            break

                    if header_idx is None or not col_map:
                        continue

                    # Parse data rows after header
                    for row in tbl[header_idx+1:]:
                        if not row:
                            continue
                        def cell(idx: int) -> str:
                            if idx is None or idx < 0 or idx >= len(row):
                                return ""
                            return (row[idx] or '').strip()

                        grantor = cell(col_map.get('grantor', -1)).replace('\n', ' ').strip()
                        grantee = cell(col_map.get('grantee', -1)).replace('\n', ' ').strip()
                        instrument = cell(col_map.get('instrument', -1)).replace('\n', ' ').strip()
                        date_cell = cell(col_map.get('date', -1))
                        rec_cell = cell(col_map.get('recording', -1))

                        # Skip empty or separator rows
                        if not any([grantor, grantee, instrument, date_cell, rec_cell]):
                            continue

                        # Extract date and book/page using strict patterns from the cell text
                        date_match = re.search(r'(\d{1,2}/\d{1,2}/\d{4})', date_cell)
                        book_match = re.search(r'([A-Z0-9]+-\d+)', rec_cell)

                        date_str = date_match.group(1) if date_match else ""
                        book_page = book_match.group(1) if book_match else rec_cell

                        if not date_str or not book_page:
                            continue

                        d = parse_date(date_str)
                        if not d:
                            continue

                        entry = ChainEntry(
                            date=d,
                            date_string=date_str,
                            grantor=grantor.upper(),
                            grantee=grantee.upper(),
                            instrument=instrument.upper(),
                            book_page=book_page,
                            is_vesting=is_vesting_deed(instrument),
                        )
                        entries.append(entry)
    except Exception as e:
        print(f"Table extraction failed: {e}")

    # Sort newest first
    entries.sort(key=lambda e: e.date, reverse=True)
    return entries

def parse_date(date_str: str) -> Optional[datetime]:
    """Parse date string in various formats."""
    if not date_str:
        return None
    
    # Clean the date string
    date_str = date_str.strip()
    
    # Try different date formats
    formats = [
        "%m/%d/%Y",  # MM/DD/YYYY
        "%m-%d-%Y",  # MM-DD-YYYY
        "%Y-%m-%d",  # YYYY-MM-DD
        "%d/%m/%Y",  # DD/MM/YYYY
    ]
    
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    # Try regex for flexible parsing
    patterns = [
        (r'(\d{1,2})/(\d{1,2})/(\d{4})', lambda g: (int(g[0]), int(g[1]), int(g[2]))),  # MM/DD/YYYY
        (r'(\d{1,2})-(\d{1,2})-(\d{4})', lambda g: (int(g[0]), int(g[1]), int(g[2]))),  # MM-DD-YYYY
        (r'(\d{4})-(\d{1,2})-(\d{1,2})', lambda g: (int(g[1]), int(g[2]), int(g[0]))),  # YYYY-MM-DD
    ]
    
    for pattern, parser in patterns:
        match = re.search(pattern, date_str)
        if match:
            try:
                month, day, year = parser(match.groups())
                return datetime(year, month, day)
            except ValueError:
                continue
    
    return None

def is_vesting_deed(instrument: str) -> bool:
    """Determine if an instrument is a vesting deed."""
    if not instrument:
        return False
    
    upper_instrument = instrument.upper().strip()
    
    # Non-vesting types (check these first)
    non_vesting = [
        'DEED OF TRUST',
        'MORTGAGE',
        'UCC',
        'ASSIGNMENT',
        'SATISFACTION',
        'RELEASE',
        'SUBORDINATION',
        'MODIFICATION',
        'LIS PENDENS',
        'AFFIDAVIT',
        'EASEMENT',
        'RIGHT OF WAY'
    ]
    
    for nv in non_vesting:
        if nv in upper_instrument:
            return False
    
    # Vesting types
    vesting = [
        'WARRANTY DEED',
        'QUITCLAIM DEED',
        'DEED'  # Generic deed last
    ]
    
    for v in vesting:
        if v in upper_instrument:
            return True
    
    return False

def parse_chain_text(text: str) -> List[ChainEntry]:
    """Parse chain of title text into structured entries using column positions."""
    lines = text.split('\n')
    entries: List[ChainEntry] = []

    # Find the header line with column positions
    header_idx = None
    col_positions: dict[str, int] = {}

    for i, line in enumerate(lines):
        # Look for the header line
        if 'GRANTOR' in line and 'GRANTEE' in line and 'INSTRUMENT' in line:
            header_idx = i
            # Get column positions from header
            col_positions = {
                'grantor': line.find('GRANTOR'),
                'grantee': line.find('GRANTEE'),
                'instrument': line.find('INSTRUMENT'),
            }
            # Date column may be labeled FILED or DATED
            filed_pos = line.find('FILED')
            dated_pos = line.find('DATED')
            col_positions['date'] = filed_pos if filed_pos != -1 else dated_pos
            # Recording column may be labeled BOOK-PAGE or RECORDING
            book_page_pos = line.find('BOOK-PAGE')
            recording_pos = line.find('RECORDING')
            col_positions['recording'] = book_page_pos if book_page_pos != -1 else recording_pos
            break

    if header_idx is None or not col_positions:
        # Fallback to regex method if no table found
        return parse_chain_text_regex_fallback(text)

    # Process table data
    in_table = True  # begin immediately after header
    current_entry_lines: List[str] = []
    has_date_in_buffer = False

    for i in range(header_idx + 1, len(lines)):
        line = lines[i]

        # End of table marker
        if '***' in line or line.strip().startswith('*'):
            if current_entry_lines:
                entry = parse_table_entry(current_entry_lines, col_positions)
                if entry:
                    entries.append(entry)
            break

        if not in_table:
            continue

        # If this line starts with a date and buffer already has a date, flush previous entry
        is_date_line = bool(re.match(r'^\s*\d{1,2}/\d{1,2}/\d{4}', line))
        if is_date_line and has_date_in_buffer and current_entry_lines:
            entry = parse_table_entry(current_entry_lines, col_positions)
            if entry:
                entries.append(entry)
            current_entry_lines = []
            has_date_in_buffer = False

        # Add line to current entry
        current_entry_lines.append(line)
        if is_date_line:
            has_date_in_buffer = True

    # Flush any remaining buffered entry
    if current_entry_lines and has_date_in_buffer:
        entry = parse_table_entry(current_entry_lines, col_positions)
        if entry:
            entries.append(entry)

    # Also augment with a simple line-wise scan to catch any entries missed by table parsing
    linewise = parse_chain_text_linewise(text)
    # De-duplicate by (date_string, instrument, book_page)
    out: List[ChainEntry] = []
    seen = set()
    for e in entries + linewise:
        key = (e.date_string, e.instrument, e.book_page)
        if key not in seen:
            seen.add(key)
            out.append(e)
    # Sort newest first for consistency
    out.sort(key=lambda e: e.date, reverse=True)
    return out

def parse_table_entry(lines: List[str], col_positions: dict) -> Optional[ChainEntry]:
    """Parse a multi-line table entry using column positions."""
    if not lines:
        return None

    # Extract data from each column
    grantor_parts: List[str] = []
    grantee_parts: List[str] = []
    instrument_parts: List[str] = []
    date_str = ""
    book_page = ""

    for line in lines:
        # Extract grantor
        if col_positions.get('grantor', -1) >= 0:
            start = col_positions['grantor']
            end = col_positions.get('grantee', len(line))
            text = line[start:end].strip()
            if text:
                grantor_parts.append(text)

        # Extract grantee
        if col_positions.get('grantee', -1) >= 0:
            start = col_positions['grantee']
            end = col_positions.get('instrument', len(line))
            text = line[start:end].strip()
            if text:
                grantee_parts.append(text)

        # Extract instrument
        if col_positions.get('instrument', -1) >= 0:
            start = col_positions['instrument']
            end = col_positions.get('recording', len(line))
            text = line[start:end].strip()
            if text:
                instrument_parts.append(text)

        # Extract date (usually only on first line)
        if not date_str and col_positions.get('date', -1) is not None and col_positions.get('date', -1) >= 0:
            start = col_positions['date']
            end = col_positions.get('recording', len(line))
            text = line[start:end].strip()
            if text and re.match(r'\d{2}/\d{2}/\d{4}', text):
                date_str = text

        # Extract recording/book-page (usually only on first line)
        if not book_page and col_positions.get('recording', -1) is not None and col_positions.get('recording', -1) >= 0:
            start = col_positions['recording']
            text = line[start:].strip()
            if text and re.match(r'[\w\d]+-[\w\d]+', text):
                book_page = text

    # Combine multi-line fields
    grantor = ' '.join(grantor_parts).strip()
    grantee = ' '.join(grantee_parts).strip()
    instrument = ' '.join(instrument_parts).strip()

    # Create entry if we have minimum required fields
    if date_str and (grantor or grantee) and book_page:
        parsed_date = parse_date(date_str)
        if parsed_date:
            return ChainEntry(
                date=parsed_date,
                date_string=date_str,
                grantor=grantor.upper(),
                grantee=grantee.upper(),
                instrument=instrument.upper(),
                book_page=book_page,
                remark="",
                is_vesting=is_vesting_deed(instrument),
                line=' '.join(lines)
            )

    return None

def preprocess_chain_text(text: str) -> str:
    """Minimal preprocessing: normalize whitespace and remove non-informational lines."""
    # Normalize Windows newlines and strip trailing spaces
    text = text.replace('\r\n', '\n')
    text = '\n'.join(line.rstrip() for line in text.split('\n'))
    return text

def parse_chain_text_regex_fallback(text: str) -> List[ChainEntry]:
    """Regex-based parsing as fallback for non-table formats."""
    text = preprocess_chain_text(text)
    lines = text.split('\n')
    entries: List[ChainEntry] = []

    for raw in lines:
        line = raw.strip()

        # Skip headers, separators, and metadata
        skip_patterns = [
            r'FILED.*GRANTOR.*GRANTEE.*INSTRUMENT',
            r'^\*+',
            r'CHAIN OF TITLE',
            r'File No\.',
            r'NAME CERTIFICATION',
            r'SELLER\s+.*\s+BUYER',
            r'OWNER:',
            r'For further information',
            r'Certified to:',
            r'New Certification Date:',
            r'By:.*',
            r'INFORMATION TO FOLLOW',
            r'^\s*$',
        ]

        if any(re.search(pattern, line, re.IGNORECASE) for pattern in skip_patterns):
            continue

        patterns = [
            r'^(\d{2}/\d{2}/\d{4})\s+(.+?)\s+(.+?)\s+((?:[\w\s]+(?:DEED|TRUST|ASSIGNMENT|MORTGAGE|UCC|SATISFACTION|RELEASE|SUBORDINATION|MODIFICATION|EXTENSION|LIS PENDENS|NOTICE|AFFIDAVIT|EASEMENT)[\w\s]*)|(?:P\s+\d+-\d+))\s+([A-Z]?\s*\d+-\d+|\w+-\w+)(?:\s+(.*))?$',
            r'^(\d{2}/\d{2}/\d{4})\s+(.+?)\s+(WARRANTY DEED|DEED OF TRUST|QUITCLAIM DEED|SPECIAL WARRANTY DEED|DEED)\s+(\d+-\d+)(?:\s+(.*))?$',
        ]

        match = None
        for pattern in patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                break

        if match:
            groups = match.groups()

            if len(groups) >= 5:
                if len(groups) == 6:
                    date_str, grantor, grantee, instrument, book_page, remark = groups
                else:
                    date_str, combined_names, instrument, book_page, remark = groups
                    name_parts = combined_names.split()
                    if len(name_parts) >= 2:
                        mid_point = len(name_parts) // 2
                        grantor = ' '.join(name_parts[:mid_point])
                        grantee = ' '.join(name_parts[mid_point:])
                    else:
                        grantor = combined_names
                        grantee = "UNKNOWN"

                parsed_date = parse_date(date_str)

                if parsed_date:
                    entry = ChainEntry(
                        date=parsed_date,
                        date_string=date_str,
                        grantor=(grantor or '').strip().upper(),
                        grantee=(grantee or '').strip().upper(),
                        instrument=instrument.strip().upper(),
                        book_page=book_page.strip(),
                        remark=(remark or '').strip(),
                        is_vesting=is_vesting_deed(instrument.strip()),
                        line=line,
                    )
                    entries.append(entry)

    return entries

def parse_chain_text_linewise(text: str) -> List[ChainEntry]:
    """Simple, robust per-line parser: match date + instrument + book-page and infer names from neighbors."""
    lines = [ln.strip() for ln in text.split('\n')]
    entries: List[ChainEntry] = []

    instr_pat = r'(SPECIAL\s+WARRANTY\s+DEED|WARRANTY\s+DEED|QUITCLAIM\s+DEED|DEED\s+OF\s+TRUST|MORTGAGE|DEED)'
    pattern = re.compile(rf'(?P<date>\d{{1,2}}/\d{{1,2}}/\d{{4}})\s+(?P<between>.*?)\s*{instr_pat}\s+(?P<book>[A-Z0-9-]+)', re.IGNORECASE)

    def valid_line(idx: int) -> bool:
        if idx < 0 or idx >= len(lines):
            return False
        ln = lines[idx]
        if not ln:
            return False
        if ln.startswith('*') or 'CHAIN OF TITLE' in ln.upper() or 'FILED GRANTOR' in ln.upper():
            return False
        return True

    for i, ln in enumerate(lines):
        m = pattern.search(ln)
        if not m:
            continue
        date_str = m.group('date')
        instrument = m.group(0)[m.start(2):m.end(0)]  # not used; use group directly below
        instrument = m.group(0)  # placeholder to avoid mypy warnings
        instrument = m.group(0)  # will be overwritten below
        # Correct instrument capture via a second regex on the matched substring to get canonical text
        sub = ln[m.start():m.end()]
        inst_m = re.search(instr_pat, sub, re.IGNORECASE)
        inst = inst_m.group(1).upper() if inst_m else ''
        book = m.group('book').strip()
        between = m.group('between').strip().upper()

        d = parse_date(date_str)
        if not d:
            continue

        # Infer grantee from the 'between' segment if present
        grantee = between if between and between not in {'', '-'} else ''

        # Infer grantor from neighbors
        grantor = ''
        if valid_line(i - 1):
            grantor = lines[i - 1].strip().upper()
        # If next line looks like continuation (e.g., suffix like INC/LLC), append to grantor when grantee already determined
        if grantor and valid_line(i + 1) and grantee:
            nxt = lines[i + 1].strip().upper()
            # Heuristic: if next line is a corporate suffix or continuation, attach to whichever of grantor/grantee seems shorter
            if any(suf in nxt for suf in [' LLC', ' INC', ' CO', ' COMPANY', ' CORP', ' TRUST', ' JR', ' SR', ' LP', ' LLP']):
                # Attach to the entity that lacks a suffix
                if not any(s in grantor for s in [' LLC', ' INC', ' CO', ' COMPANY', ' CORP', ' TRUST', ' JR', ' SR', ' LP', ' LLP']):
                    grantor = (grantor + ' ' + nxt).strip()
                elif not any(s in grantee for s in [' LLC', ' INC', ' CO', ' COMPANY', ' CORP', ' TRUST', ' JR', ' SR', ' LP', ' LLP']):
                    grantee = (grantee + ' ' + nxt).strip()

        # Fallbacks
        if not grantee and valid_line(i + 1):
            grantee = lines[i + 1].strip().upper()

        if not grantor:
            grantor = 'UNKNOWN'
        if not grantee:
            grantee = 'UNKNOWN'

        entry = ChainEntry(
            date=d,
            date_string=date_str,
            grantor=grantor,
            grantee=grantee,
            instrument=inst,
            book_page=book,
            is_vesting=is_vesting_deed(inst),
        )
        entries.append(entry)

    # Sort newest first
    entries.sort(key=lambda e: e.date, reverse=True)
    return entries

def parse_chain_text_by_labels(text: str) -> List[ChainEntry]:
    """Parse entries using GRANTOR/GRANTEE/INSTRUMENT/DATED/RECORDING labels in flattened text."""
    entries: List[ChainEntry] = []
    # Collapse whitespace to make regex simpler
    flat = re.sub(r"\s+", " ", text)

    # Pattern tries to capture label-driven rows
    pattern = re.compile(
        r"GRANTOR\s*:?\s*(?P<grantor>.+?)\s+GRANTEE\s*:?\s*(?P<grantee>.+?)\s+"
        r"(?P<instrument>SPECIAL\s+WARRANTY\s+DEED|WARRANTY\s+DEED|QUITCLAIM\s+DEED|DEED\s+OF\s+TRUST|MORTGAGE|DEED)\s+"
        r"DATED\s*:?\s*(?P<date>\d{1,2}/\d{1,2}/\d{4})\s+RECORDING\s*:?\s*(?P<record>\d{3,6}-\d{1,7})",
        re.IGNORECASE,
    )

    for m in pattern.finditer(flat):
        date_str = m.group("date")
        parsed_date = parse_date(date_str)
        if not parsed_date:
            continue
        grantor = m.group("grantor").strip().upper()
        grantee = m.group("grantee").strip().upper()
        instrument = m.group("instrument").strip().upper()
        record = m.group("record").strip()
        entry = ChainEntry(
            date=parsed_date,
            date_string=date_str,
            grantor=grantor,
            grantee=grantee,
            instrument=instrument,
            book_page=record,
            is_vesting=is_vesting_deed(instrument),
        )
        entries.append(entry)

    # Sort newest first
    entries.sort(key=lambda e: e.date, reverse=True)
    return entries

def parse_single_entry(line: str, all_lines: list, current_idx: int) -> Optional[ChainEntry]:
    """Parse a single chain entry from the current line."""
    # Extract date
    date_match = re.match(r'^(\d{1,2}/\d{1,2}/\d{4})\s+(.+)', line)
    if not date_match:
        return None
    
    date_str = date_match.group(1)
    parsed_date = parse_date(date_str)
    if not parsed_date:
        return None
    
    # Get the rest of the line after the date
    rest = date_match.group(2).strip()
    
    # Check if entry continues on next line(s)
    combined_text = rest
    next_idx = current_idx + 1
    while next_idx < len(all_lines):
        next_line = all_lines[next_idx].strip()
        # Stop if we hit another date, empty line, or separator
        if not next_line or re.match(r'^\d{2}/\d{2}/\d{4}', next_line) or '***' in next_line:
            break
        # Add continuation lines
        combined_text += ' ' + next_line
        next_idx += 1
    
    # Parse the combined text for grantor, grantee, instrument, and book-page
    # Pattern: GRANTOR [text] GRANTEE [text] INSTRUMENT_TYPE BOOK-PAGE
    
    # Find book-page anywhere in the combined text (pattern: digits-digits)
    book_page = ""
    text_without_book = combined_text
    book_page_matches = list(re.finditer(r'(\d{3,6}-\d{1,7})', combined_text))
    if book_page_matches:
        last_match = book_page_matches[-1]
        book_page = last_match.group(1)
        text_without_book = combined_text[: last_match.start()].strip()
    
    # Find instrument type (working backwards from book-page)
    instrument = ""
    instrument_patterns = [
        r'(SPECIAL\s+WARRANTY\s+DEED|WARRANTY\s+DEED|QUITCLAIM\s+DEED|DEED\s+OF\s+TRUST|MORTGAGE|UCC\s+FINANCING\s+STATEMENT|DEED)\s*$',
        r'(SPECIAL\s+WARRANTY\s+DEED|WARRANTY\s+DEED|QUITCLAIM\s+DEED|DEED\s+OF\s+TRUST|MORTGAGE|UCC\s+FINANCING\s+STATEMENT|DEED)\b',
    ]
    
    for pattern in instrument_patterns:
        inst_match = re.search(pattern, text_without_book, re.IGNORECASE)
        if inst_match:
            instrument = inst_match.group(1).strip()
            text_without_inst = text_without_book[:inst_match.start()].strip()
            break
    
    if not instrument:
        # If instrument couldn't be found, attempt a heuristic: look for the word DEED
        inst_match = re.search(r'\b([A-Z ]*DEED)\b', text_without_book)
        if inst_match:
            instrument = inst_match.group(1).strip()
        else:
            return None
    
    # Split remaining text into grantor and grantee
    # This is the tricky part - use heuristics
    grantor, grantee = split_grantor_grantee(text_without_inst)
    
    if not grantor or not grantee:
        return None
    
    return ChainEntry(
        date=parsed_date,
        date_string=date_str,
        grantor=grantor.upper(),
        grantee=grantee.upper(),
        instrument=instrument.upper(),
        book_page=book_page,
        is_vesting=is_vesting_deed(instrument),
        line=line
    )

def split_grantor_grantee(text: str) -> tuple[str, str]:
    """Split text into grantor and grantee using heuristics."""
    # Clean up the text
    text = ' '.join(text.split())  # Normalize whitespace
    
    # Common patterns that indicate end of grantor/start of grantee
    split_patterns = [
        # Look for repeated words (often company names are repeated)
        r'(.*?)\s+(LLC|INC|CORP|CORPORATION|COMPANY)\s+(.*?\2.*)',
        # Look for "TO" separator
        r'(.*?)\s+TO\s+(.*)',
        # Look for double space or large gap
        r'(.*?)\s{2,}(.*)',
    ]
    
    for pattern in split_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            grantor = match.group(1).strip()
            grantee = match.groups()[-1].strip()
            if grantor and grantee:
                return grantor, grantee
    
    # Fallback: split at halfway point
    words = text.split()
    if len(words) >= 2:
        mid = len(words) // 2
        grantor = ' '.join(words[:mid])
        grantee = ' '.join(words[mid:])
        return grantor, grantee
    
    return "", ""

def parse_table_data(table_rows: List[List[str]]) -> List[ChainEntry]:
    """Parse table data from a list of rows."""
    entries = []
    
    # Skip header row
    for row in table_rows[1:]:
        if len(row) >= 5:
            date_str = row[3] if len(row) > 3 else ""
            parsed_date = parse_date(date_str)
            
            if parsed_date:
                entry = ChainEntry(
                    date=parsed_date,
                    date_string=date_str,
                    grantor=row[0].upper(),
                    grantee=row[1].upper(),
                    instrument=row[2].upper(),
                    book_page=row[4] if len(row) > 4 else "",
                    is_vesting=is_vesting_deed(row[2])
                )
                entries.append(entry)
    
    return entries

def get_24_month_chain(entries: List[ChainEntry], processing_date: datetime = None) -> List[ChainEntry]:
    """Return the minimum set of vesting deeds that cover at least the last 24 months.

    Rules:
    - Always consider vesting deeds only.
    - Start with the most recent vesting deed.
    - If the time from the most recent deed to processing_date is >= 24 months, return it alone.
    - Otherwise, include earlier vesting deeds until the earliest included deed is at least 24 months before processing_date.
    - If no vesting deeds exist, return an empty list.
    """
    if processing_date is None:
        processing_date = datetime.now()

    # Filter for vesting deeds and sort by date descending (newest first)
    vesting_deeds = [e for e in entries if e.is_vesting]
    vesting_deeds.sort(key=lambda x: x.date, reverse=True)

    if not vesting_deeds:
        return []

    twenty_four_months_ago = processing_date - timedelta(days=730)

    # Start with most recent vesting deed
    selected: List[ChainEntry] = [vesting_deeds[0]]

    # If coverage from newest deed to now already exceeds 24 months, done
    if selected[0].date <= twenty_four_months_ago:
        return selected

    # Otherwise, add older deeds until coverage threshold is reached
    for deed in vesting_deeds[1:]:
        selected.append(deed)
        # Earliest included deed is the last in selected (since we append in descending order)
        earliest_included = selected[-1].date
        if earliest_included <= twenty_four_months_ago:
            break

    # Ensure newest-first ordering
    selected.sort(key=lambda x: x.date, reverse=True)
    return selected

def process_title_document(file_path: Optional[str] = None, file_bytes: Optional[bytes] = None, output_path: str = None, template_path: str = None) -> tuple[bool, str, List[ChainEntry], List[ChainEntry]]:
    """Process title document (PDF/DOCX) from path or bytes and optionally create output."""
    if not file_path and not file_bytes:
        return False, "Either file_path or file_bytes must be provided.", [], []

    try:
        file_ext = ''
        if file_path:
            file_ext = os.path.splitext(file_path)[1].lower()
        else:
            # Assume PDF from bytes, as it's the output of the splitter
            file_ext = '.pdf'
        
        text = ""
        entries = []
        
        if file_ext == '.pdf':
            # Use a context manager for both BytesIO and file open
            stream_manager = io.BytesIO(file_bytes) if file_bytes else open(file_path, 'rb')
            with stream_manager as stream:
                try:
                    entries = extract_table_entries_from_pdf(stream)
                    stream.seek(0)  # Reset stream for re-reading
                except Exception:
                    entries = [] # Ensure entries is empty on failure
                
                text = extract_text_from_pdf(stream)

        elif file_ext == '.docx':
            doc_source = io.BytesIO(file_bytes) if file_bytes else file_path
            doc = Document(doc_source)
            text = '\n'.join([p.text for p in doc.paragraphs])
        else:
            return False, f"Unsupported file type: {file_ext}", [], []
        
        if not text.strip() and not entries:
            return False, "No text extracted from document", [], []
        
        # If table-based extraction was unsuccessful, fall back to text parsing
        if not entries:
            entries = parse_chain_text(text)
        
        if not entries:
            return False, "No chain entries found", [], []
        
        # Get 24-month chain
        chain_deeds = get_24_month_chain(entries)
        
        # Create output document if requested
        if output_path and template_path:
            success = create_title_document(chain_deeds, output_path, template_path)
            if not success:
                return False, "Failed to create output document", chain_deeds, entries
        
        msg = f"Found {len(entries)} total entries, {len(chain_deeds)} vesting deeds in 24-month chain"
        return True, msg, chain_deeds, entries
        
    except Exception as e:
        return False, f"Error: {str(e)}", [], []

def create_title_document(chain_deeds: List[ChainEntry], output_path: str, template_path: str) -> bool:
    """Create the output document from template."""
    try:
        doc = Document(template_path)
        
        # Find the chain table
        chain_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                header_text = ' '.join([cell.text.upper() for cell in table.rows[0].cells])
                if 'GRANTOR' in header_text and 'GRANTEE' in header_text:
                    chain_table = table
                    break
        
        if chain_table:
            # Clear existing data rows
            while len(chain_table.rows) > 1:
                chain_table._element.remove(chain_table.rows[-1]._element)
            
            # Add chain data
            for deed in chain_deeds:
                row = chain_table.add_row()
                cells = row.cells
                if len(cells) >= 5:
                    cells[0].text = deed.grantor
                    cells[1].text = deed.grantee
                    cells[2].text = deed.instrument
                    cells[3].text = deed.date_string
                    cells[4].text = deed.book_page
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating document: {e}")
        return False