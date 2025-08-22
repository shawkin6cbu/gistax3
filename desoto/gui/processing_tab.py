import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinterdnd2 import DND_FILES
from desoto.services.document_splitter import process_comprehensive_document
from desoto.services.title_chain import process_title_document
from desoto.services.tax_document import process_tax_document
from docx.text.paragraph import Paragraph
import threading
import os
from docx import Document
import re
import subprocess
import sys

class ProcessingTab(ttk.Frame):
    def __init__(self, parent, shared_data):
        super().__init__(parent, padding=20)
        self.shared_data = shared_data

        # Main layout grid
        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=1)

        # --- LEFT COLUMN ---
        left_frame = ttk.Frame(self)
        left_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        left_frame.columnconfigure(0, weight=1)

        # --- Document Drop Zone (Single unified drop zone) ---
        doc_drop_frame = ttk.LabelFrame(left_frame, text="Title Search Document", padding=15)
        doc_drop_frame.grid(row=0, column=0, sticky="ew", pady=(0, 20))
        doc_drop_frame.columnconfigure(1, weight=1)
        
        self.search_doc_var = tk.StringVar()
        self.search_doc_entry = self._create_file_drop_row(
            doc_drop_frame, 
            "Drop Title Search:", 
            self.search_doc_var, 
            self.browse_search_document, 
            self._drop_on_search_doc
        )
        
        # Processing status
        self.process_status_var = tk.StringVar(value="No document loaded")
        status_label = ttk.Label(doc_drop_frame, textvariable=self.process_status_var, wraplength=350)
        status_label.grid(row=1, column=0, columnspan=3, sticky="w", pady=(10, 0))
        
        # Progress bar for document processing (styled for visual appeal)
        self.doc_progress_frame = ttk.Frame(doc_drop_frame)
        self.doc_progress_frame.grid(row=2, column=0, columnspan=3, sticky="ew", pady=(5, 0))
        self.doc_progress_frame.columnconfigure(0, weight=1)
        
        # Create styled progress bar with percentage label
        progress_container = ttk.Frame(self.doc_progress_frame)
        progress_container.grid(row=0, column=0, sticky="ew")
        progress_container.columnconfigure(0, weight=1)
        
        self.doc_progress = ttk.Progressbar(
            progress_container, 
            mode="determinate",
            style="Processing.Horizontal.TProgressbar"
        )
        self.doc_progress.grid(row=0, column=0, sticky="ew", pady=(2, 0))
        
        self.doc_progress_label = ttk.Label(
            progress_container, 
            text="",
            font=("Segoe UI", 9)
        )
        self.doc_progress_label.grid(row=1, column=0, sticky="w", pady=(2, 0))
        
        # Initially hide the progress bar
        self.doc_progress_frame.grid_remove()
        
        # Configure progress bar style for visual appeal
        style = ttk.Style()
        style.configure(
            "Processing.Horizontal.TProgressbar",
            thickness=20,
            troughcolor='#e0e0e0',
            background='#4CAF50',
            darkcolor='#45a049',
            lightcolor='#66BB6A',
            bordercolor='#e0e0e0',
            relief='flat'
        )
        

        # --- Property Information ---
        prop_frame = ttk.LabelFrame(left_frame, text="Property Information", padding=15)
        prop_frame.grid(row=1, column=0, sticky="ew", pady=(0, 20))
        prop_frame.columnconfigure(1, weight=1)
        
        self.pin_var = tk.StringVar()
        self.address_var = tk.StringVar()
        self.owner_var = tk.StringVar()
        self.city_var = tk.StringVar()
        self.legal_desc_var = tk.StringVar()
        
        self._create_entry_row(prop_frame, "Parcel:", self.pin_var, 0)
        self._create_entry_row(prop_frame, "Address:", self.address_var, 1)
        self._create_entry_row(prop_frame, "Owner:", self.owner_var, 2)
        self._create_entry_row(prop_frame, "City/State/ZIP:", self.city_var, 3)
        self._create_entry_row(prop_frame, "Legal Desc:", self.legal_desc_var, 4)

        # --- Tax Information ---
        tax_frame = ttk.LabelFrame(left_frame, text="Tax Information", padding=15)
        tax_frame.grid(row=2, column=0, sticky="ew", pady=(0, 20))
        tax_frame.columnconfigure(1, weight=1)
        
        self.tax_2024_total_var = tk.StringVar()
        self.tax_2024_paid_var = tk.StringVar(value="PAID")
        self.tax_2024_date_paid_var = tk.StringVar()
        self.tax_2025_est_var = tk.StringVar()
        
        self._create_entry_row(tax_frame, "2024 Total:", self.tax_2024_total_var, 0)
        self._create_entry_row(tax_frame, "Date Paid:", self.tax_2024_date_paid_var, 1)
        self._create_entry_row(tax_frame, "2024 Status:", self.tax_2024_paid_var, 2, is_combo=True)
        self._create_entry_row(tax_frame, "2025 Estimated:", self.tax_2025_est_var, 3)

        # --- RIGHT COLUMN ---
        right_frame = ttk.Frame(self)
        right_frame.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        right_frame.columnconfigure(0, weight=1)

        # --- Title Chain Summary ---
        title_frame = ttk.LabelFrame(right_frame, text="Title Chain Summary", padding=15)
        title_frame.grid(row=0, column=0, sticky="ew", pady=(0, 20))
        title_frame.columnconfigure(0, weight=1)
        
        self.title_summary_var = tk.StringVar(value="Not loaded yet.")
        ttk.Label(title_frame, textvariable=self.title_summary_var, wraplength=400, justify="left").grid(row=0, column=0, columnspan=2, sticky="w")
        ttk.Button(title_frame, text="View & Edit Title Chain", command=self.view_title_details).grid(row=1, column=0, sticky="w", pady=(10,0))
        
        # --- Document Details ---
        doc_details_frame = ttk.LabelFrame(right_frame, text="Document Details", padding=15)
        doc_details_frame.grid(row=1, column=0, sticky="ew", pady=(0, 20))
        doc_details_frame.columnconfigure(1, weight=1)
        self.lender_var = tk.StringVar()
        self.borrower_var = tk.StringVar()
        self._create_entry_row(doc_details_frame, "Lender:", self.lender_var, 0)
        self._create_entry_row(doc_details_frame, "Borrower:", self.borrower_var, 1)

        # --- Document Generation ---
        doc_gen_frame = ttk.LabelFrame(right_frame, text="Document Generation", padding=15)
        doc_gen_frame.grid(row=2, column=0, sticky="ew")
        doc_gen_frame.columnconfigure(1, weight=1)
        
        self.output_path_var = tk.StringVar()
        self._create_entry_row(doc_gen_frame, "Output Path:", self.output_path_var, 0, browse_btn=True)
        
        self.generate_btn = ttk.Button(doc_gen_frame, text="Generate Document", command=self.generate_document)
        self.generate_btn.grid(row=1, column=1, sticky="w", pady=(10,0))
        
        self.progress = ttk.Progressbar(doc_gen_frame, mode="indeterminate")
        self.progress.grid(row=2, column=0, columnspan=3, sticky="ew", pady=(10,0))

    def _create_entry_row(self, parent, label_text, var, row, browse_btn=False, is_combo=False):
        ttk.Label(parent, text=label_text).grid(row=row, column=0, sticky="e", padx=(0, 10), pady=5)
        
        if is_combo:
            combo = ttk.Combobox(
                parent,
                textvariable=var,
                values=["PAID", "UNPAID", "PARTIAL"],
                state="normal",
                width=28
            )
            combo.grid(row=row, column=1, sticky="ew", pady=5)
        else:
            entry = ttk.Entry(parent, textvariable=var, width=40)
            entry.grid(row=row, column=1, sticky="ew", pady=5, padx=(0, 5 if browse_btn else 0))
        
        if browse_btn:
            ttk.Button(parent, text="...", width=4, command=self.browse_output).grid(row=row, column=2, pady=5)
        
        parent.columnconfigure(1, weight=1)

    def _create_file_drop_row(self, parent, label_text, var, browse_cmd, drop_cmd):
        parent.drop_target_register(DND_FILES)
        parent.dnd_bind("<<Drop>>", drop_cmd)

        ttk.Label(parent, text=label_text).grid(row=0, column=0, sticky="w", pady=5)
        
        entry = ttk.Entry(parent, textvariable=var, width=30)
        entry.grid(row=0, column=1, sticky="ew", pady=5, padx=(10, 5))
        
        ttk.Button(parent, text="Browse", command=browse_cmd).grid(row=0, column=2, sticky="e", pady=5)

        # Auto-process when file path changes
        var.trace_add('write', lambda *_: self.process_document_if_valid(var))
        return entry

    def process_document_if_valid(self, var):
        path = (var.get() or "").strip()
        if path and os.path.exists(path) and path.lower().endswith(('.pdf', '.docx')):
            self.process_search_document()

    def _drop_on_search_doc(self, event=None):
        if event:
            path = event.data.strip('{}')
            if path.lower().endswith(('.pdf', '.docx')):
                self.search_doc_var.set(path)
        self.process_search_document()

    def browse_search_document(self):
        path = filedialog.askopenfilename(
            title="Select Title Search Document", 
            filetypes=[("PDF files","*.pdf"), ("Word documents", "*.docx"), ("All files","*.*")]
        )
        if path:
            self.search_doc_var.set(path)

    def show_progress(self, message="Processing...", progress=0):
        """Show and update the progress bar with a message."""
        self.doc_progress_frame.grid()
        self.doc_progress['value'] = progress
        self.doc_progress_label['text'] = f"{message} ({progress}%)"
        self.update_idletasks()

    def hide_progress(self):
        """Hide the progress bar."""
        self.doc_progress_frame.grid_remove()
        self.doc_progress['value'] = 0
        self.doc_progress_label['text'] = ""

    def process_search_document(self):
        """Process the comprehensive search document, extracting both chain and tax info."""
        file_path = self.search_doc_var.get().strip()
        if not file_path or not os.path.exists(file_path):
            return
        
        self.process_status_var.set("Processing document...")
        self.title_summary_var.set("Processing...")
        self.tax_2024_total_var.set("Processing...")
        self.tax_2024_date_paid_var.set("Processing...")
        
        threading.Thread(target=self._process_search_document_thread, args=(file_path,), daemon=True).start()

    def _process_search_document_thread(self, file_path):
        try:
            # Update progress
            self.after(0, lambda: self.show_progress("Reading document", 20))
            
            # Process the comprehensive document
            success, msg, results = process_comprehensive_document(file_path)
            
            # Update progress
            self.after(0, lambda: self.show_progress("Extracting data", 60))
            
            def update_ui():
                self.show_progress("Finalizing results", 90)
                
                self.process_status_var.set(results.get('status', msg))
                
                # Update chain information
                chain_entries = results.get('chain_entries', [])
                all_entries = results.get('all_entries', [])
                
                if chain_entries:
                    self.shared_data.set_data("title_chain_kept", chain_entries)
                    self.shared_data.set_data("title_chain_all", all_entries)
                    summary = f"Found {len(chain_entries)} vesting deeds in 24-month chain."
                    self.title_summary_var.set(summary)
                else:
                    self.shared_data.set_data("title_chain_kept", [])
                    self.shared_data.set_data("title_chain_all", [])
                    self.title_summary_var.set("No chain data found.")
                
                # Update tax information
                tax_total = results.get('tax_total')
                tax_date_paid = results.get('tax_date_paid')
                
                if tax_total or tax_date_paid:
                    self.tax_2024_total_var.set(tax_total or "")
                    self.tax_2024_date_paid_var.set(tax_date_paid or "")
                    if tax_date_paid:
                        self.tax_2024_paid_var.set("PAID")
                    
                    self.shared_data.update_data({
                        "tax_2024_total": tax_total or "",
                        "tax_2024_date_paid": tax_date_paid or "",
                        "tax_2024_paid_status": "PAID" if tax_date_paid else self.tax_2024_paid_var.get()
                    })
                else:
                    self.tax_2024_total_var.set("")
                    self.tax_2024_date_paid_var.set("")
                
                # Complete progress
                self.show_progress("Processing complete!", 100)
                
                # Update status with results summary
                status_parts = []
                if chain_entries:
                    status_parts.append(f"{len(chain_entries)} vesting deeds found")
                if tax_total:
                    status_parts.append(f"Tax: ${tax_total}")
                if not status_parts:
                    status_parts.append("No data extracted")
                
                final_status = "✓ " + " | ".join(status_parts) if success else "⚠ " + msg
                self.process_status_var.set(final_status)
                
                # Hide progress bar after a short delay
                self.after(1500, self.hide_progress)
                    
            self.after(0, update_ui)
            
        except Exception as e:
            error_msg = f"Error processing document: {str(e)}"
            self.after(0, lambda: self.process_status_var.set("Processing failed"))
            self.after(0, lambda: self.hide_progress())
            self.after(0, lambda: messagebox.showerror("Processing Error", error_msg))

    def load_from_tabs(self):
        self.pin_var.set(self.shared_data.get_data("parcel_pin") or "")
        self.address_var.set(self.shared_data.get_data("parcel_address") or "")
        self.owner_var.set(self.shared_data.get_data("parcel_owner") or "")
        self.city_var.set(self.shared_data.get_data("parcel_city_state_zip") or "")
        self.legal_desc_var.set(self.shared_data.get_data("parcel_legal_description") or "")
        
        self.tax_2024_total_var.set(self.shared_data.get_data("tax_2024_total") or "")
        self.tax_2024_paid_var.set(self.shared_data.get_data("tax_2024_paid_status") or "PAID")
        self.tax_2024_date_paid_var.set(self.shared_data.get_data("tax_2024_date_paid") or "")
        self.tax_2025_est_var.set(self.shared_data.get_data("tax_2025_estimated") or "")

        results = self.shared_data.get_data("title_chain_kept")
        self.title_summary_var.set(f"{len(results)} vesting deeds found." if results else "No title chain data found.")

    def view_title_details(self):
        all_entries = self.shared_data.get_data("title_chain_all") or []
        kept_entries = self.shared_data.get_data("title_chain_kept") or []

        if not all_entries:
            messagebox.showinfo("Title Chain Details", "No title chain data has been loaded.")
            return

        details_win = tk.Toplevel(self)
        details_win.title("Manage Title Chain Entries")
        details_win.geometry("900x700")

        paned_window = ttk.PanedWindow(details_win, orient=tk.VERTICAL)
        paned_window.pack(fill="both", expand=True, padx=10, pady=10)

        cols = ("Date", "Grantor", "Grantee", "Instrument", "Book-Page")
        sort_states = {'keep': {'col': 'Date', 'rev': True}, 'other': {'col': 'Date', 'rev': True}}

        def sort_tree(tree, tree_key, col):
            state = sort_states[tree_key]
            reverse = not state['rev'] if state['col'] == col else False
            state.update({'col': col, 'rev': reverse})

            items = [(tree.set(child, col), child) for child in tree.get_children('')]
            
            if col == 'Date':
                from datetime import datetime
                items.sort(key=lambda it: datetime.strptime(it[0], "%m/%d/%Y"), reverse=reverse)
            elif col == 'Book-Page':
                items.sort(key=lambda it: [int(p) for p in re.match(r'(\d+)-(\d+)', it[0]).groups()] if re.match(r'(\d+)-(\d+)', it[0]) else [0,0], reverse=reverse)
            else:
                items.sort(key=lambda it: it[0].lower(), reverse=reverse)

            for index, (_, child) in enumerate(items):
                tree.move(child, '', index)

            for c in cols:
                arrow = ' ↓' if reverse else ' ↑'
                tree.heading(c, text=c + (arrow if c == col else ''))

        def create_tree_view(parent, key):
            frame = ttk.LabelFrame(parent, text="Kept" if key == 'keep' else "Other Entries", padding=10)
            parent.add(frame, weight=1 if key == 'keep' else 2)
            tree = ttk.Treeview(frame, columns=cols, show="headings")
            for col in cols:
                tree.heading(col, text=col, command=lambda c=col, t=tree, k=key: sort_tree(t, k, c))
                tree.column(col, width=150, anchor="w")
            tree.pack(fill="both", expand=True)
            return tree

        keep_tree = create_tree_view(paned_window, 'keep')
        other_tree = create_tree_view(paned_window, 'other')

        all_entries_map = {(e.date_string, e.book_page, e.instrument): e for e in all_entries}
        
        def populate_trees():
            kept_set = {(e.date_string, e.book_page, e.instrument) for e in kept_entries}
            
            for tree in [keep_tree, other_tree]:
                for item in tree.get_children():
                    tree.delete(item)

            for entry in sorted(all_entries, key=lambda e: e.date, reverse=True):
                values = (entry.date_string, entry.grantor, entry.grantee, entry.instrument, entry.book_page)
                key = (entry.date_string, entry.book_page, entry.instrument)
                (keep_tree if key in kept_set else other_tree).insert("", "end", values=values)

            sort_tree(keep_tree, 'keep', 'Date')
            sort_tree(other_tree, 'other', 'Date')

        def move_item(from_tree, to_tree, event):
            selected_item = from_tree.selection()
            if not selected_item: return
            item_values = from_tree.item(selected_item, "values")
            to_tree.insert("", "end", values=item_values)
            from_tree.delete(selected_item)

        keep_tree.bind("<Double-1>", lambda e: move_item(keep_tree, other_tree, e))
        other_tree.bind("<Double-1>", lambda e: move_item(other_tree, keep_tree, e))

        def on_close():
            new_kept_keys = {tuple(keep_tree.item(child, "values")) for child in keep_tree.get_children()}
            new_kept_entries = [all_entries_map.get((v[0], v[4], v[3])) for v in new_kept_keys]
            new_kept_entries = [e for e in new_kept_entries if e]
            new_kept_entries.sort(key=lambda e: e.date, reverse=True)
            
            self.shared_data.set_data("title_chain_kept", new_kept_entries)
            self.title_summary_var.set(f"{len(new_kept_entries)} vesting deeds in chain.")
            details_win.destroy()

        details_win.protocol("WM_DELETE_WINDOW", on_close)
        populate_trees()

    def browse_output(self):
        path = filedialog.asksaveasfilename(
            title="Save Document As",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
        )
        if path:
            self.output_path_var.set(path)

    def generate_document(self):
        self.sync_to_shared_data()
        doc_path = (self.search_doc_var.get() or "").strip()
        out_path = self.output_path_var.get().strip()

        # --- MODIFICATION: Find a unique filename ---
        # 1. Determine base path
        output_path = out_path or os.path.abspath("TitleDocs.docx")
        if not out_path and doc_path:
            try:
                abs_doc_path = os.path.abspath(doc_path)
                grandparent_dir = os.path.dirname(os.path.dirname(abs_doc_path))
                output_path = os.path.join(grandparent_dir, "TitleDocs.docx")
            except Exception:
                try:
                    output_path = os.path.join(os.path.dirname(os.path.abspath(doc_path)), "TitleDocs.docx")
                except Exception:
                    pass
        
        # 2. Check for existence and find a unique name by appending a number
        if os.path.exists(output_path):
            directory, filename = os.path.split(output_path)
            base, ext = os.path.splitext(filename)
            
            # Strip existing numbers from the base to avoid names like "TitleDocs23.docx"
            match = re.search(r'(\d+)$', base)
            if match:
                base = base[:match.start()]

            counter = 2
            while True:
                new_filename = f"{base}{counter}{ext}"
                new_path = os.path.join(directory, new_filename)
                if not os.path.exists(new_path):
                    output_path = new_path  # Found a unique name
                    break
                counter += 1
        
        # Update the entry box with the final path
        self.output_path_var.set(output_path)
        # --- END MODIFICATION ---

        self.generate_btn.config(state="disabled")
        self.progress.start()
        threading.Thread(target=self._generate_document_thread, args=(output_path,), daemon=True).start()

    def _open_file_os(self, path):
        """Opens a file with the default OS application in a cross-platform way."""
        try:
            abs_path = os.path.abspath(path)
            if sys.platform == "win32":
                os.startfile(abs_path)
            elif sys.platform == "darwin":  # macOS
                subprocess.call(['open', abs_path])
            else:  # Linux and other Unix-like systems
                subprocess.call(['xdg-open', abs_path])
            return True, f"Opening document: {abs_path}"
        except Exception as e:
            return False, f"Could not open file: {e}"

    def _generate_document_thread(self, output_path):
        try:
            success, msg = self._create_full_document(output_path)
            
            final_msg = msg
            if success:
                open_success, open_msg = self._open_file_os(output_path)
                if not open_success:
                    final_msg += f"\n\nNote: {open_msg}"
            
            self.after(0, lambda: messagebox.showinfo("Success", final_msg) if success else messagebox.showerror("Error", msg))

        except Exception as e:
            error_msg = f"An unexpected error occurred: {str(e)}"
            import traceback
            traceback.print_exc()  # Print full traceback to console
            self.after(0, lambda: messagebox.showerror("Error", error_msg))
        finally:
            self.after(0, self.progress.stop)
            self.after(0, lambda: self.generate_btn.config(state="normal"))

    def get_template_path(self):
        try:
            base_dir = os.path.dirname(os.path.dirname((os.path.dirname(os.path.abspath(__file__)))))
            path = os.path.join(base_dir, "templates", "td_tmplt2.docx")
            return path if os.path.exists(path) else None
        except Exception:
            return None

    def sync_to_shared_data(self):
        self.shared_data.update_data({
            "parcel_pin": self.pin_var.get(),
            "parcel_address": self.address_var.get(),
            "parcel_owner": self.owner_var.get(),
            "parcel_city_state_zip": self.city_var.get(),
            "parcel_legal_description": self.legal_desc_var.get(),
            "tax_2024_total": self.tax_2024_total_var.get(),
            "tax_2024_paid_status": self.tax_2024_paid_var.get(),
            "tax_2024_date_paid": self.tax_2024_date_paid_var.get(),
            "tax_2025_estimated": self.tax_2025_est_var.get(),
            "lender": self.lender_var.get(),
            "borrower": self.borrower_var.get(),
        })

    def _create_full_document(self, output_path):
        """Create the final document with proper placeholder mapping."""
        template_path = self.get_template_path()
        if not template_path:
            return False, "Template file 'td_tmplt2.docx' not found."
        
        doc = Document(template_path)

        def smart_title_case(text):
            if not text or text.strip().isdigit() or '$' in text:
                return text
            preserve_upper = ['LLC', 'PLLC', 'INC', 'CO', 'CORP', 'LP', 'LLP', 'PA', 'PC', 'LTD', 'II', 'III', 'IV', 'JR', 'SR', 'MS', 'US']
            words = text.split()
            result = []
            for word in words:
                clean_word = word.strip('.,;:')
                if clean_word.upper() in preserve_upper:
                    result.append(clean_word.upper() + word[len(clean_word):])
                else:
                    result.append(word.capitalize())
            return ' '.join(result)

        values_map = {
            "PARCEL": self.pin_var.get(),
            "PROPSTRE": smart_title_case(self.address_var.get()),
            "SLRLAST": smart_title_case(self.owner_var.get()),
            "CITY_STATE_ZIP": smart_title_case(self.city_var.get()),
            "LEGAL_DESC": smart_title_case(self.legal_desc_var.get()),
            "TAXAMT": ("$" + self.tax_2024_total_var.get()) if self.tax_2024_total_var.get() else "",
            "TAXDAT": self.tax_2024_date_paid_var.get(),
            "TAX_2025_EST": ("$" + self.tax_2025_est_var.get()) if self.tax_2025_est_var.get() else "",
            "Lender": smart_title_case(self.lender_var.get()),
            "BYRLAST": smart_title_case(self.borrower_var.get()),
            "LOAN_AMOUNT": "",
        }

        print("--- FORMAT-PRESERVING FIX: Values being mapped ---")
        for key, value in values_map.items():
            print(f"  - {{{key}}}: '{value}'")

        # --- NEW, FORMAT-PRESERVING REPLACEMENT LOGIC ---
        def replace_text_in_element(element, values):
            if isinstance(element, Paragraph):
                paragraphs = [element]
            else:
                paragraphs = element.paragraphs
                
            for p in paragraphs:
                for key, value in values.items():
                    placeholder = f"{{{key}}}"
                    while True:
                        full_text = "".join(run.text for run in p.runs)
                        if placeholder not in full_text:
                            break

                        start_char_idx = full_text.find(placeholder)
                        end_char_idx = start_char_idx + len(placeholder)

                        char_count = 0
                        start_run_info, end_run_info = None, None
                        
                        runs_to_process = list(p.runs)
                        for i, run in enumerate(runs_to_process):
                            run_len = len(run.text)
                            if not start_run_info and start_char_idx < char_count + run_len:
                                start_run_info = {'index': i, 'pos': start_char_idx - char_count}
                            if not end_run_info and end_char_idx <= char_count + run_len:
                                end_run_info = {'index': i, 'pos': end_char_idx - char_count}
                                break
                            char_count += run_len

                        if start_run_info and end_run_info:
                            start_idx = start_run_info['index']
                            start_pos = start_run_info['pos']
                            end_idx = end_run_info['index']
                            end_pos = end_run_info['pos']
                            
                            start_run = runs_to_process[start_idx]
                            
                            if start_idx == end_idx:
                                before = start_run.text[:start_pos]
                                after = start_run.text[end_pos:]
                                start_run.text = f"{before}{value}{after}"
                            else:
                                start_run.text = start_run.text[:start_pos] + str(value)
                                for i in range(start_idx + 1, end_idx):
                                    runs_to_process[i].text = ""
                                end_run = runs_to_process[end_idx]
                                end_run.text = end_run.text[end_pos:]

        # Apply the replacement to all paragraphs and tables
        for paragraph in doc.paragraphs:
            replace_text_in_element(paragraph, values_map)

        for table in doc.tables:
            header_text = ' '.join([cell.text.upper() for cell in table.rows[0].cells])
            if 'GRANTOR' in header_text and 'GRANTEE' in header_text:
                continue

            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_element(cell, values_map)

        # --- Fill Title Chain Table ---
        chain_deeds = self.shared_data.get_data("title_chain_kept") or []
        chain_table = None
        for table in doc.tables:
            header_text = ' '.join([cell.text.upper() for cell in table.rows[0].cells])
            if 'GRANTOR' in header_text and 'GRANTEE' in header_text:
                chain_table = table
                break

        if chain_table and chain_deeds:
            while len(chain_table.rows) > 1:
                chain_table._element.remove(chain_table.rows[-1]._element)
            for deed in sorted(chain_deeds, key=lambda d: d.date):
                row = chain_table.add_row().cells
                row[0].text = deed.grantor
                row[1].text = deed.grantee
                row[2].text = deed.instrument
                row[3].text = deed.date_string
                row[4].text = deed.book_page
            print(f"Added {len(chain_deeds)} deeds to the title chain table.")

        doc.save(output_path)
        return True, f"Document successfully generated at:\n{output_path}"
