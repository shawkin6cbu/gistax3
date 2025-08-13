import os, sys, re
from docx import Document


def main():
    template_path = os.path.join('gistax','templates','td_tmplt2.docx')
    print('TEMPLATE', template_path)
    if not template_path or not os.path.exists(template_path):
        print('NO_TEMPLATE')
        return
    doc = Document(template_path)

    # Collect all text nodes
    all_texts = []
    for p in doc.paragraphs:
        all_texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_texts.append(p.text)

    # Dump unique placeholder-like tokens
    combined = '\n'.join(all_texts)
    print('DOC_TEXT_LEN', len(combined))

    # Heuristics: braces, ALLCAPS tokens, or known names
    tokens = set()
    brace_pat = re.compile(r'\{[^}]{2,40}\}')
    for m in brace_pat.finditer(combined):
        tokens.add(m.group(0))
    caps_pat = re.compile(r'\b[A-Z]{3,}\b')
    for m in caps_pat.finditer(combined):
        if any(prefix in m.group(0) for prefix in ['PARCEL', 'PROP', 'SLR', 'BYR', 'TAX', 'LEND']):
            tokens.add(m.group(0))

    print('TOKENS_FOUND', sorted(tokens))
    # Print lines containing target fragments to see exact spelling
    targets = ['PARCEL', 'PROP', 'SLR', 'BYR', 'TAX', 'Lender', 'LENDER']
    for line in combined.split('\n'):
        if any(t in line for t in targets) or '{' in line:
            print('LINE:', line)

if __name__ == '__main__':
    main()

