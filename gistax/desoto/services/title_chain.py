import re
import PyPDF2
from datetime import datetime, timedelta
from dataclasses import dataclass
from typing import List, Optional
from docx import Document
import os

@dataclass
class ChainEntry:
    date: datetime
    date_string: str
    grantor: str
    grantee: str
    instrument: str
    book_page: str
    remark: str = ""
    is_vesting: bool = False
    line: str = ""

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF - simplified to use PyPDF2 only."""
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF extraction failed: {e}")
    return text

def parse_date(date_str: str) -> Optional[datetime]:
    """Parse date string in various formats."""
    if not date_str:
        return None
    
    # Clean the date string
    date_str = date_str.strip()
    
    # Try different date formats
    formats = [
        "%m/%d/%Y",  # MM/DD/YYYY
        "%m-%d-%Y",  # MM-DD-YYYY
        "%Y-%m-%d",  # YYYY-MM-DD
        "%d/%m/%Y",  # DD/MM/YYYY
    ]
    
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    # Try regex for flexible parsing
    patterns = [
        (r'(\d{1,2})/(\d{1,2})/(\d{4})', lambda m: (int(m[1]), int(m[2]), int(m[3]))),  # MM/DD/YYYY
        (r'(\d{1,2})-(\d{1,2})-(\d{4})', lambda m: (int(m[1]), int(m[2]), int(m[3]))),  # MM-DD-YYYY
        (r'(\d{4})-(\d{1,2})-(\d{1,2})', lambda m: (int(m[2]), int(m[3]), int(m[1]))),  # YYYY-MM-DD
    ]
    
    for pattern, parser in patterns:
        match = re.search(pattern, date_str)
        if match:
            try:
                month, day, year = parser(match.groups())
                return datetime(year, month, day)
            except ValueError:
                continue
    
    return None

def is_vesting_deed(instrument: str) -> bool:
    """Determine if an instrument is a vesting deed."""
    if not instrument:
        return False
    
    upper_instrument = instrument.upper().strip()
    
    # Non-vesting types (check these first)
    non_vesting = [
        'DEED OF TRUST',
        'MORTGAGE',
        'UCC',
        'ASSIGNMENT',
        'SATISFACTION',
        'RELEASE',
        'SUBORDINATION',
        'MODIFICATION',
        'LIS PENDENS',
        'AFFIDAVIT',
        'EASEMENT',
        'RIGHT OF WAY'
    ]
    
    for nv in non_vesting:
        if nv in upper_instrument:
            return False
    
    # Vesting types
    vesting = [
        'WARRANTY DEED',
        'QUITCLAIM DEED',
        'DEED'  # Generic deed last
    ]
    
    for v in vesting:
        if v in upper_instrument:
            return True
    
    return False

def parse_chain_text(text: str) -> List[ChainEntry]:
    """Main parser - simplified to handle your PDF format."""
    entries = []
    lines = text.split('\n')
    
    # Find the start of the chain data
    start_idx = -1
    for i, line in enumerate(lines):
        if 'FILED' in line and 'GRANTOR' in line and 'GRANTEE' in line:
            start_idx = i + 1  # Start after header
            break
    
    if start_idx == -1:
        print("No chain header found")
        return entries
    
    # Process entries
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines and markers
        if not line or '***' in line or 'NAME CERTIFICATION' in line:
            i += 1
            continue
        
        # Check if line starts with a date
        date_match = re.match(r'^(\d{2}/\d{2}/\d{4})', line)
        if date_match:
            # This is the start of an entry
            entry = parse_single_entry(line, lines, i)
            if entry:
                entries.append(entry)
        
        i += 1
    
    return entries

def parse_single_entry(line: str, all_lines: list, current_idx: int) -> Optional[ChainEntry]:
    """Parse a single chain entry from the current line."""
    # Extract date
    date_match = re.match(r'^(\d{2}/\d{2}/\d{4})\s+(.+)', line)
    if not date_match:
        return None
    
    date_str = date_match.group(1)
    parsed_date = parse_date(date_str)
    if not parsed_date:
        return None
    
    # Get the rest of the line after the date
    rest = date_match.group(2).strip()
    
    # Check if entry continues on next line(s)
    combined_text = rest
    next_idx = current_idx + 1
    while next_idx < len(all_lines):
        next_line = all_lines[next_idx].strip()
        # Stop if we hit another date, empty line, or separator
        if not next_line or re.match(r'^\d{2}/\d{2}/\d{4}', next_line) or '***' in next_line:
            break
        # Add continuation lines
        combined_text += ' ' + next_line
        next_idx += 1
    
    # Parse the combined text for grantor, grantee, instrument, and book-page
    # Pattern: GRANTOR [text] GRANTEE [text] INSTRUMENT_TYPE BOOK-PAGE
    
    # Find book-page at the end (pattern: digits-digits)
    book_page_match = re.search(r'(\d+-\d+)\s*$', combined_text)
    if not book_page_match:
        return None
    
    book_page = book_page_match.group(1)
    text_without_book = combined_text[:book_page_match.start()].strip()
    
    # Find instrument type (working backwards from book-page)
    instrument = ""
    instrument_patterns = [
        r'(WARRANTY DEED|DEED OF TRUST|QUITCLAIM DEED|MORTGAGE|UCC FINANCING STATEMENT|DEED)\s*$',
    ]
    
    for pattern in instrument_patterns:
        inst_match = re.search(pattern, text_without_book, re.IGNORECASE)
        if inst_match:
            instrument = inst_match.group(1).strip()
            text_without_inst = text_without_book[:inst_match.start()].strip()
            break
    
    if not instrument:
        return None
    
    # Split remaining text into grantor and grantee
    # This is the tricky part - use heuristics
    grantor, grantee = split_grantor_grantee(text_without_inst)
    
    if not grantor or not grantee:
        return None
    
    return ChainEntry(
        date=parsed_date,
        date_string=date_str,
        grantor=grantor.upper(),
        grantee=grantee.upper(),
        instrument=instrument.upper(),
        book_page=book_page,
        is_vesting=is_vesting_deed(instrument),
        line=line
    )

def split_grantor_grantee(text: str) -> tuple[str, str]:
    """Split text into grantor and grantee using heuristics."""
    # Clean up the text
    text = ' '.join(text.split())  # Normalize whitespace
    
    # Common patterns that indicate end of grantor/start of grantee
    split_patterns = [
        # Look for repeated words (often company names are repeated)
        r'(.*?)\s+(LLC|INC|CORP|CORPORATION|COMPANY)\s+(.*?\2.*)',
        # Look for "TO" separator
        r'(.*?)\s+TO\s+(.*)',
        # Look for double space or large gap
        r'(.*?)\s{2,}(.*)',
    ]
    
    for pattern in split_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            grantor = match.group(1).strip()
            grantee = match.groups()[-1].strip()
            if grantor and grantee:
                return grantor, grantee
    
    # Fallback: split at halfway point
    words = text.split()
    if len(words) >= 2:
        mid = len(words) // 2
        grantor = ' '.join(words[:mid])
        grantee = ' '.join(words[mid:])
        return grantor, grantee
    
    return "", ""

def parse_table_data(table_rows: List[List[str]]) -> List[ChainEntry]:
    """Parse table data from a list of rows."""
    entries = []
    
    # Skip header row
    for row in table_rows[1:]:
        if len(row) >= 5:
            date_str = row[3] if len(row) > 3 else ""
            parsed_date = parse_date(date_str)
            
            if parsed_date:
                entry = ChainEntry(
                    date=parsed_date,
                    date_string=date_str,
                    grantor=row[0].upper(),
                    grantee=row[1].upper(),
                    instrument=row[2].upper(),
                    book_page=row[4] if len(row) > 4 else "",
                    is_vesting=is_vesting_deed(row[2])
                )
                entries.append(entry)
    
    return entries

def get_24_month_chain(entries: List[ChainEntry], processing_date: datetime = None) -> List[ChainEntry]:
    """Get vesting deeds covering at least 24 months."""
    if processing_date is None:
        processing_date = datetime.now()
    
    cutoff_date = processing_date - timedelta(days=730)  # 24 months
    
    # Filter for vesting deeds only
    vesting_deeds = [e for e in entries if e.is_vesting]
    vesting_deeds.sort(key=lambda x: x.date, reverse=True)
    
    if not vesting_deeds:
        return []
    
    result = []
    
    # Include all vesting deeds within 24 months
    for deed in vesting_deeds:
        if deed.date >= cutoff_date:
            result.append(deed)
    
    # If no deeds within 24 months, include the most recent one
    if not result and vesting_deeds:
        result.append(vesting_deeds[0])
    
    # Sort by date (newest first)
    result.sort(key=lambda x: x.date, reverse=True)
    return result

def process_title_document(file_path: str, output_path: str = None, template_path: str = None) -> tuple[bool, str, List[ChainEntry]]:
    """Process title document and optionally create output."""
    try:
        # Extract text based on file type
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            doc = Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs])
        else:
            return False, f"Unsupported file type: {file_ext}", []
        
        if not text.strip():
            return False, "No text extracted from document", []
        
        # Parse the chain entries
        entries = parse_chain_text(text)
        
        if not entries:
            return False, "No chain entries found", []
        
        # Get 24-month chain
        chain_deeds = get_24_month_chain(entries)
        
        # Create output document if requested
        if output_path and template_path:
            success = create_title_document(chain_deeds, output_path, template_path)
            if not success:
                return False, "Failed to create output document", chain_deeds
        
        msg = f"Found {len(entries)} total entries, {len(chain_deeds)} vesting deeds in 24-month chain"
        return True, msg, chain_deeds
        
    except Exception as e:
        return False, f"Error: {str(e)}", []

def create_title_document(chain_deeds: List[ChainEntry], output_path: str, template_path: str) -> bool:
    """Create the output document from template."""
    try:
        doc = Document(template_path)
        
        # Find the chain table
        chain_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                header_text = ' '.join([cell.text.upper() for cell in table.rows[0].cells])
                if 'GRANTOR' in header_text and 'GRANTEE' in header_text:
                    chain_table = table
                    break
        
        if chain_table:
            # Clear existing data rows
            while len(chain_table.rows) > 1:
                chain_table._element.remove(chain_table.rows[-1]._element)
            
            # Add chain data
            for deed in chain_deeds:
                row = chain_table.add_row()
                cells = row.cells
                if len(cells) >= 5:
                    cells[0].text = deed.grantor
                    cells[1].text = deed.grantee
                    cells[2].text = deed.instrument
                    cells[3].text = deed.date_string
                    cells[4].text = deed.book_page
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating document: {e}")
        return False