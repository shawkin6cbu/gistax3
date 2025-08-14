import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinterdnd2 import DND_FILES
from desoto.services.title_chain import (
    extract_text_from_pdf, 
    parse_chain_text, 
    get_24_month_chain,
    process_title_document
)
import threading
import os
from docx import Document

class ProcessingTab(ttk.Frame):
    def __init__(self, parent, shared_data):
        super().__init__(parent, padding=10)
        self.shared_data = shared_data

        # Main layout frames
        left_frame = ttk.Frame(self)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 5))
        right_frame = ttk.Frame(self)
        right_frame.pack(side="right", fill="both", expand=True, padx=(5, 0))

        # --- Property Information ---
        prop_frame = ttk.LabelFrame(left_frame, text="Property Information", padding=10)
        prop_frame.pack(fill="x", pady=(0, 10))
        self.pin_var = tk.StringVar()
        self.address_var = tk.StringVar()
        self.owner_var = tk.StringVar()
        self.city_var = tk.StringVar()
        self.legal_desc_var = tk.StringVar()
        self._create_entry_row(prop_frame, "PIN:", self.pin_var, 0)
        self._create_entry_row(prop_frame, "Address:", self.address_var, 1)
        self._create_entry_row(prop_frame, "Owner:", self.owner_var, 2)
        self._create_entry_row(prop_frame, "City/State/ZIP:", self.city_var, 3)
        self._create_entry_row(prop_frame, "Legal Desc:", self.legal_desc_var, 4)

        # --- Tax Information ---
        tax_frame = ttk.LabelFrame(left_frame, text="Tax Information", padding=10)
        tax_frame.pack(fill="x", pady=(0, 10))
        self.tax_2024_total_var = tk.StringVar()
        self.tax_2024_paid_var = tk.StringVar()
        self.tax_2024_date_paid_var = tk.StringVar()
        self.tax_2025_est_var = tk.StringVar()
        self._create_entry_row(tax_frame, "2024 Total:", self.tax_2024_total_var, 0)
        self._create_entry_row(tax_frame, "Date Paid:", self.tax_2024_date_paid_var, 1)
        
        # Paid status dropdown (moved to row 2)
        self._create_entry_row(tax_frame, "2024 Status:", self.tax_2024_paid_var, 2)
        paid_frame = ttk.Frame(tax_frame)
        paid_frame.grid(row=2, column=1, sticky="ew", pady=2)
        self.paid_combo = ttk.Combobox(
            paid_frame,
            textvariable=self.tax_2024_paid_var,
            values=["PAID", "UNPAID", "PARTIAL"],
            state="normal",
            width=28,
        )
        self.paid_combo.pack(side="left")
        
        self._create_entry_row(tax_frame, "2025 Estimated:", self.tax_2025_est_var, 3)

        # --- Title Chain Document ---
        doc_frame = ttk.LabelFrame(left_frame, text="Drag and drop extracted title chain PDF", padding=10)
        doc_frame.pack(fill="x", pady=(0, 10))
        self.title_doc_var = tk.StringVar()
        doc_row = ttk.Frame(doc_frame)
        doc_row.pack(fill="x")
        # Styled drop area
        drop_style = ttk.Style()
        drop_style.configure("DropArea.TEntry", relief="solid", borderwidth=1)

        ttk.Label(doc_row, text="File:").pack(side="left")
        self.title_doc_entry = ttk.Entry(doc_row, textvariable=self.title_doc_var, width=40, style="DropArea.TEntry")
        self.title_doc_entry.pack(side="left", fill="x", expand=True, padx=(6,6))
        ttk.Button(doc_row, text="Browse", command=self.browse_title_document).pack(side="left")

        # Drop handler
        def _drop_on_doc(event):
            path = event.data.strip('{}')
            if path.lower().endswith(('.pdf', '.docx')):
                self.title_doc_var.set(path)
                self.process_title_document()
        doc_frame.drop_target_register(DND_FILES)
        doc_frame.dnd_bind("<<Drop>>", _drop_on_doc)
        
        # Auto-process when file path is set via typing/paste/browse
        def on_title_path_change(*_):
            path = (self.title_doc_var.get() or "").strip()
            if path and os.path.exists(path) and path.lower().endswith(('.pdf', '.docx')):
                self.process_title_document()
        try:
            self.title_doc_var.trace_add('write', lambda *_: on_title_path_change())  # type: ignore[attr-defined]
        except Exception:
            self.title_doc_var.trace('w', lambda *_: on_title_path_change())

        # --- Title Chain Summary ---
        title_frame = ttk.LabelFrame(left_frame, text="Title Chain Summary", padding=10)
        title_frame.pack(fill="x", pady=(0, 10))
        self.title_summary_var = tk.StringVar(value="Not loaded yet.")
        ttk.Label(title_frame, textvariable=self.title_summary_var, wraplength=250).pack(anchor="w")
        ttk.Button(title_frame, text="View Details", command=self.view_title_details).pack(pady=(5,0))

        # --- Document Details ---
        doc_details_frame = ttk.LabelFrame(right_frame, text="Document Details", padding=10)
        doc_details_frame.pack(fill="x", pady=(0, 10))
        self.lender_var = tk.StringVar()
        self.borrower_var = tk.StringVar()
        # removed loan amount, writer, date, notes per request
        self._create_entry_row(doc_details_frame, "Lender:", self.lender_var, 0)
        self._create_entry_row(doc_details_frame, "Borrower:", self.borrower_var, 1)
        # No Loan Amount field


        # --- Document Generation ---
        doc_gen_frame = ttk.LabelFrame(right_frame, text="Document Generation", padding=10)
        doc_gen_frame.pack(fill="x", pady=(0, 10))
        self.output_path_var = tk.StringVar()
        self._create_entry_row(doc_gen_frame, "Output Path:", self.output_path_var, 0, browse_btn=True)
        self.generate_btn = ttk.Button(doc_gen_frame, text="Generate Document", command=self.generate_document)
        self.generate_btn.grid(row=1, column=1, sticky="w", pady=(10,0))
        self.progress = ttk.Progressbar(doc_gen_frame, mode="indeterminate")
        self.progress.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(5,0))

        # Auto-loading is handled by other tabs; no manual button needed

    def browse_title_document(self):
        path = filedialog.askopenfilename(
            title="Select Title Chain Document", 
            filetypes=[("PDF files","*.pdf"), ("Word documents", "*.docx"), ("All files","*.*")]
        )
        if path:
            self.title_doc_var.set(path)

    def process_title_document(self):
        """Process the title document using improved extraction."""
        file_path = self.title_doc_var.get().strip()
        if not file_path or not os.path.exists(file_path):
            messagebox.showerror("Error", "Please select a valid title document file.")
            return
        
        self.title_summary_var.set("Processing document...")
        threading.Thread(target=self._process_title_document_thread, args=(file_path,), daemon=True).start()

    def _process_title_document_thread(self, file_path):
        """Thread worker for processing title documents."""
        try:
            print(f"Processing title document: {file_path}")
            
            # Use the improved process_title_document function
            success, msg, chain_deeds, all_entries = process_title_document(file_path)
            
            if success:
                # Store both the initial "kept" list and the full list
                self.shared_data.set_data("title_chain_kept", chain_deeds)
                self.shared_data.set_data("title_chain_all", all_entries)
                
                # Update summary text
                summary = f"Found {len(chain_deeds)} vesting deeds in 24-month chain."
                self.after(0, lambda: self.title_summary_var.set(summary))
                print(msg)
            else:
                # On failure, clear out any old data
                self.shared_data.set_data("title_chain_kept", [])
                self.shared_data.set_data("title_chain_all", [])
                self.after(0, lambda: self.title_summary_var.set(msg))
                self.after(0, lambda: messagebox.showerror("Processing Error", msg))
            
        except Exception as e:
            error_msg = f"Error processing document: {str(e)}"
            print(f"Title processing error: {error_msg}")
            import traceback
            traceback.print_exc()
            self.after(0, lambda: self.title_summary_var.set("Processing failed."))
            self.after(0, lambda: messagebox.showerror("Processing Error", error_msg))

    def _create_entry_row(self, parent, label_text, var, row, browse_btn=False):
        ttk.Label(parent, text=label_text).grid(row=row, column=0, sticky="e", padx=(0, 5), pady=2)
        if label_text != "2024 Status:":  # Skip for the paid status combo
            entry = ttk.Entry(parent, textvariable=var, width=30)
            entry.grid(row=row, column=1, sticky="ew", padx=(0, 5 if browse_btn else 0), pady=2)
        if browse_btn:
            ttk.Button(parent, text="...", width=3, command=self.browse_output).grid(row=row, column=2, pady=2)
        parent.columnconfigure(1, weight=1)

    def load_from_tabs(self):
        """Load data from other tabs into this tab."""
        # Auto-sync from SharedData
        self.pin_var.set(self.shared_data.get_data("parcel_pin") or "")
        self.address_var.set(self.shared_data.get_data("parcel_address") or "")
        self.owner_var.set(self.shared_data.get_data("parcel_owner") or "")
        self.city_var.set(self.shared_data.get_data("parcel_city_state_zip") or "")
        self.legal_desc_var.set(self.shared_data.get_data("parcel_legal_description") or "")
        
        # Tax data
        tax_total = self.shared_data.get_data("tax_2024_total") or ""
        self.tax_2024_total_var.set(tax_total)
        self.tax_2024_paid_var.set(self.shared_data.get_data("tax_2024_paid_status") or "PAID")
        self.tax_2024_date_paid_var.set(self.shared_data.get_data("tax_2024_date_paid") or "")
        self.tax_2025_est_var.set(self.shared_data.get_data("tax_2025_estimated") or "")

        # Title chain
        results = self.shared_data.get_data("title_chain_kept")
        if results:
            self.title_summary_var.set(f"{len(results)} vesting deeds found.")
            print(f"Loaded {len(results)} title chain results")
        else:
            self.title_summary_var.set("No title chain data found.")

    def view_title_details(self):
        all_entries = self.shared_data.get_data("title_chain_all") or []
        kept_entries = self.shared_data.get_data("title_chain_kept") or []

        if not all_entries:
            messagebox.showinfo("Title Chain Details", "No title chain data has been loaded.")
            return

        details_win = tk.Toplevel(self)
        details_win.title("Manage Title Chain Entries")
        details_win.geometry("900x700")

        # Use a paned window to make sections resizable
        paned_window = ttk.PanedWindow(details_win, orient=tk.VERTICAL)
        paned_window.pack(fill="both", expand=True, padx=10, pady=10)

        # --- Keep Section ---
        keep_frame = ttk.LabelFrame(paned_window, text="Keep", padding=10)
        paned_window.add(keep_frame, weight=1)

        cols = ("Date", "Grantor", "Grantee", "Instrument", "Book-Page")
        keep_tree = ttk.Treeview(keep_frame, columns=cols, show="headings")
        for col in cols:
            keep_tree.heading(col, text=col)
            keep_tree.column(col, width=150, anchor="w")
        keep_tree.pack(fill="both", expand=True)

        # --- Other Section ---
        other_frame = ttk.LabelFrame(paned_window, text="Other Entries", padding=10)
        paned_window.add(other_frame, weight=2)

        other_tree = ttk.Treeview(other_frame, columns=cols, show="headings")
        for col in cols:
            other_tree.heading(col, text=col)
            other_tree.column(col, width=150, anchor="w")
        other_tree.pack(fill="both", expand=True)

        # --- Data Handling ---
        # Create a dictionary for quick lookups of all entries by a unique key
        all_entries_map = {(e.date_string, e.book_page, e.instrument): e for e in all_entries}
        
        # Keep track of which items are in which tree
        kept_keys = set()
        other_keys = set()

        def populate_trees():
            # Clear existing items
            for item in keep_tree.get_children():
                keep_tree.delete(item)
            for item in other_tree.get_children():
                other_tree.delete(item)
            
            kept_set = {(e.date_string, e.book_page, e.instrument) for e in kept_entries}
            
            # Sort all entries by date
            sorted_entries = sorted(all_entries, key=lambda e: e.date, reverse=True)

            for entry in sorted_entries:
                key = (entry.date_string, entry.book_page, entry.instrument)
                values = (entry.date_string, entry.grantor, entry.grantee, entry.instrument, entry.book_page)
                if key in kept_set:
                    keep_tree.insert("", "end", values=values, iid=f"kept_{len(kept_keys)}")
                    kept_keys.add(key)
                else:
                    other_tree.insert("", "end", values=values, iid=f"other_{len(other_keys)}")
                    other_keys.add(key)

        def move_to_other(event):
            selected_item = keep_tree.selection()
            if not selected_item:
                return
            
            item_values = keep_tree.item(selected_item, "values")
            key = (item_values[0], item_values[4], item_values[3]) # date, book-page, instrument

            if key in kept_keys:
                kept_keys.remove(key)
                other_keys.add(key)
                
                other_tree.insert("", "end", values=item_values, iid=f"other_{len(other_keys)}")
                keep_tree.delete(selected_item)

        def move_to_keep(event):
            selected_item = other_tree.selection()
            if not selected_item:
                return

            item_values = other_tree.item(selected_item, "values")
            key = (item_values[0], item_values[4], item_values[3])

            if key in other_keys:
                other_keys.remove(key)
                kept_keys.add(key)

                keep_tree.insert("", "end", values=item_values, iid=f"kept_{len(kept_keys)}")
                other_tree.delete(selected_item)
        
        keep_tree.bind("<Double-1>", move_to_other)
        other_tree.bind("<Double-1>", move_to_keep)

        def on_close():
            # When the window is closed, update the shared data
            new_kept_entries = [all_entries_map[key] for key in kept_keys if key in all_entries_map]
            # Sort by date before saving
            new_kept_entries.sort(key=lambda e: e.date, reverse=True)
            self.shared_data.set_data("title_chain_kept", new_kept_entries)
            
            # Update the summary on the main tab
            self.title_summary_var.set(f"{len(new_kept_entries)} vesting deeds in chain.")
            print(f"Updated kept deeds to {len(new_kept_entries)} entries.")
            
            details_win.destroy()

        details_win.protocol("WM_DELETE_WINDOW", on_close)
        
        # Initial population
        populate_trees()

    def browse_output(self):
        path = filedialog.asksaveasfilename(
            title="Save Document As",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
        )
        if path:
            self.output_path_var.set(path)

    def generate_document(self):
        #Sync current values to shared data
        self.sync_to_shared_data()
        
        # Determine output path
        doc_path = (self.title_doc_var.get() or "").strip()
        out_path = self.output_path_var.get().strip()

        if doc_path:
            try:
                gp = os.path.abspath(os.path.join(os.path.dirname(doc_path), os.pardir))
                output_path = os.path.join(gp, "TitleDocs.docx")
            except Exception:
                output_path = out_path or os.path.abspath("TitleDocs.docx")
        else:
            output_path = out_path or os.path.abspath("TitleDocs.docx")

        self.generate_btn.config(state="disabled")
        self.progress.start()

        threading.Thread(target=self._generate_document_thread, args=(output_path,), daemon=True).start()

    def _generate_document_thread(self, output_path):
        try:
            success, msg = self._create_full_document(output_path)
            if success:
                self.after(0, lambda: messagebox.showinfo("Success", msg))
            else:
                self.after(0, lambda: messagebox.showerror("Error", msg))
        except Exception as e:
            error_msg = f"An unexpected error has occured: {str(e)}"
            import traceback
            traceback.print_exc()
            self.after(0, lambda: messagebox.showerror("Error", error_msg))
        finally:
            self.after(0, self.progress.stop)
            self.after(0, lambda: self.generate_btn.config(state="normal"))

    def get_template_path(self):
        try:
            current_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            path = os.path.join(current_dir, "templates", "td_tmplt2.docx")
            return path if os.path.exists(path) else None
        except Exception:
            return None

    def sync_to_shared_data(self):
        """Sync all Processing tab fields to shared data."""
        self.shared_data.update_data({
            "parcel_pin": self.pin_var.get(),
            "parcel_address": self.address_var.get(),
            "parcel_owner": self.owner_var.get(),
            "parcel_city_state_zip": self.city_var.get(),
            "parcel_legal_description": self.legal_desc_var.get(),
            "tax_2024_total": self.tax_2024_total_var.get(),
            "tax_2024_paid_status": self.tax_2024_paid_var.get(),
            "tax_2024_date_paid": self.tax_2024_date_paid_var.get(),
            "tax_2025_estimated": self.tax_2025_est_var.get(),
            "lender": self.lender_var.get(),
            "borrower": self.borrower_var.get(),
        })

    def _create_full_document(self, output_path):
        """Create the final document with proper placeholder mapping."""
        template_path = self.get_template_path()
        if not template_path:
            return False, "Template file 'td_tmplt2.docx' not found."

        # Process title document if not already done
        doc_path = (self.title_doc_var.get() or "").strip()
        if doc_path and os.path.exists(doc_path):
            current_results = self.shared_data.get_data("title_chain_kept")
            if not current_results:
                print("Processing title document for generation...")
                from desoto.services.title_chain import process_title_document
                success, msg, chain_deeds, all_entries = process_title_document(doc_path)
                if success:
                    self.shared_data.set_data("title_chain_kept", chain_deeds)
                    self.shared_data.set_data("title_chain_all", all_entries)

        doc = Document(template_path)

        def smart_title_case(text):
            """Convert to title case but preserve common business abbreviations."""
            if not text or text.strip().isdigit() or '$' in text:
                return text  # Skip numbers and dollar amounts
            
            # Abbreviations to keep uppercase
            preserve_upper = ['LLC', 'PLLC', 'INC', 'CO', 'CORP', 'LP', 'LLP', 'PA', 'PC', 'LTD', 'II', 'III', 'IV', 'JR', 'SR', 'MS', 'US']
            
            words = text.split()
            result = []
            for word in words:
                # Check if word (without punctuation) should be preserved
                clean_word = word.strip('.,;:')
                if clean_word.upper() in preserve_upper:
                    result.append(clean_word.upper() + word[len(clean_word):])  # Keep punctuation
                else:
                    # Title case the word
                    result.append(word.capitalize())
            
            return ' '.join(result)

        # Build values map
        values_map = {
            "PARCEL": self.pin_var.get(),
            "PROPSTRE": smart_title_case(self.address_var.get()),
            "SLRLAST": smart_title_case(self.owner_var.get()),
            "CITY_STATE_ZIP": self.city_var.get(),
            "LEGAL_DESC": self.legal_desc_var.get(),
            "TAXAMT": ("$" + self.tax_2024_total_var.get()) if self.tax_2024_total_var.get() else "",
            "TAXDAT": self.tax_2024_date_paid_var.get(),
            "TAX_2025_EST": ("$" + self.tax_2025_est_var.get()) if self.tax_2025_est_var.get() else "",
            "Lender": smart_title_case(self.lender_var.get()),
            "BYRLAST": smart_title_case(self.borrower_var.get()),
            "LOAN_AMOUNT": "",
        }

        print("VALUES TO REPLACE:")
        for key, value in values_map.items():
            print(f"  {key}: '{value}'")

        # Simple replacement function
        def replace_text(text):
            for key, value in values_map.items():
                # Replace {KEY} format
                text = text.replace(f"{{{key}}}", value)
                # Also replace bare KEY format  
                text = text.replace(key, value)
                # Handle dollar sign variants
                text = text.replace(f"${{{key}}}", value)
            return text

        # Replace in all paragraphs
        def replace_in_paragraph(paragraph):
            for run in paragraph.runs:
                if run.text:
                    for key, value in values_map.items():
                        # Replace {KEY} format
                        run.text = run.text.replace(f"{{{key}}}", value)
                        # Handle dollar sign variants
                        run.text = run.text.replace(f"${{{key}}}", value)
                        # Don't replace bare KEY format - too risky for headers

        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        # Replace in all tables (but skip chain table)
        for table in doc.tables:
            # Check if this is the chain table
            is_chain_table = False
            if len(table.rows) > 0:
                header_text = ' '.join([cell.text.upper() for cell in table.rows[0].cells])
                if 'GRANTOR' in header_text and 'GRANTEE' in header_text:
                    is_chain_table = True
            
            if not is_chain_table:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)

        # Fill Title Chain Table
        chain_deeds = self.shared_data.get_data("title_chain_kept") or []
        
        # Find the chain table
        chain_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                header_text = ' '.join([cell.text.upper() for cell in table.rows[0].cells])
                if 'GRANTOR' in header_text and 'GRANTEE' in header_text:
                    chain_table = table
                    break

        if chain_table and chain_deeds:
            # Clear existing data rows (keep header)
            while len(chain_table.rows) > 1:
                chain_table._element.remove(chain_table.rows[-1]._element)

            # Add chain data
            for deed in chain_deeds:
                row = chain_table.add_row()
                cells = row.cells
                if len(cells) >= 5:
                    cells[0].text = deed.grantor
                    cells[1].text = deed.grantee
                    cells[2].text = deed.instrument
                    cells[3].text = deed.date_string
                    cells[4].text = deed.book_page

            print(f"Added {len(chain_deeds)} deeds to chain table")

        doc.save(output_path)
        return True, f"Document successfully generated at:\n{output_path}"