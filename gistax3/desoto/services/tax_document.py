import re
import PyPDF2
from typing import Optional, Tuple
from docx import Document

def extract_tax_info_from_pdf(pdf_path: str) -> Tuple[bool, str, Optional[str], Optional[str]]:
    """
    Extract 2024 tax information from a tax document PDF.
    
    Returns:
        Tuple of (success, message, total_amount, date_paid)
    """
    try:
        # Try pdfplumber first for better table extraction
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    # Try to extract tables first
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table:
                                if row:
                                    text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                    # Also get regular text
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}, falling back to PyPDF2")
            text = ""
        
        # Fallback to PyPDF2 if pdfplumber fails
        if not text.strip():
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                return False, f"Failed to extract text from PDF: {str(e)}", None, None
        
        if not text.strip():
            return False, "No text extracted from document", None, None
        
        # Parse the extracted text for 2024 tax information
        total_amount, date_paid = parse_tax_text(text)
        
        if total_amount or date_paid:
            msg = "Successfully extracted tax information"
            return True, msg, total_amount, date_paid
        else:
            return False, "Could not find 2024 tax information in document", None, None
            
    except Exception as e:
        return False, f"Error processing tax document: {str(e)}", None, None

def parse_tax_text(text: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Parse tax document text to extract 2024 tax total and date paid.
    
    Returns:
        Tuple of (total_amount, date_paid)
    """
    lines = text.split('\n')
    total_amount = None
    date_paid = None
    
    # Patterns to look for
    year_2024_pattern = r'2024'
    
    # Look for 2024 row in various formats
    for i, line in enumerate(lines):
        if '2024' in line:
            # Clean up the line
            clean_line = ' '.join(line.split())
            
            # Pattern 1: Table format with pipes or tabs
            # Example: "2024 | $3,177.00 | $149.74 | PAID 01/29/2025 | $321.91"
            if '|' in clean_line or '\t' in clean_line:
                parts = re.split(r'[|\t]', clean_line)
                # Look for PAID status with date
                for j, part in enumerate(parts):
                    if 'PAID' in part:
                        date_match = re.search(r'PAID\s+(\d{1,2}/\d{1,2}/\d{4})', part)
                        if date_match:
                            date_paid = date_match.group(1)
                        # The total is usually the last monetary value in the row
                        # Look for the rightmost dollar amount
                        money_values = re.findall(r'\$?([\d,]+\.?\d*)', clean_line)
                        if money_values:
                            # Get the last value (usually the total)
                            total_amount = money_values[-1].replace(',', '')
                            # Clean up - remove trailing periods if decimal is .00
                            if '.' not in total_amount:
                                total_amount = total_amount
                            break
            
            # Pattern 2: Space-separated format
            # Example: "2024 $3,177.00 $149.74 PAID 01/29/2025 $321.91"
            else:
                # Look for PAID with date
                paid_match = re.search(r'PAID\s+(\d{1,2}/\d{1,2}/\d{4})', clean_line)
                if paid_match:
                    date_paid = paid_match.group(1)
                
                # Find all dollar amounts in the line
                money_pattern = r'\$?([\d,]+\.?\d*)'
                money_matches = re.findall(money_pattern, clean_line)
                
                # Filter to only valid monetary amounts
                valid_amounts = []
                for amount in money_matches:
                    # Skip year-like numbers (4 digits without decimal)
                    if len(amount) == 4 and '.' not in amount:
                        continue
                    # Must have at least 2 digits or a decimal
                    if '.' in amount or len(amount.replace(',', '')) >= 2:
                        valid_amounts.append(amount.replace(',', ''))
                
                # The total is typically the last amount on the line
                if valid_amounts:
                    total_amount = valid_amounts[-1]
            
            # Also check the next few lines for continuation
            if not date_paid and i < len(lines) - 1:
                next_line = lines[i + 1]
                date_match = re.search(r'(\d{1,2}/\d{1,2}/\d{4})', next_line)
                if date_match:
                    date_paid = date_match.group(1)
            
            # If we found the data, stop searching
            if total_amount or date_paid:
                break
    
    # Alternative pattern: Look for explicit TOTAL or STATUS sections
    if not total_amount or not date_paid:
        # Combine all lines for easier searching
        full_text = ' '.join(lines)
        
        # Look for patterns like "2024...TOTAL...321.91"
        total_pattern = r'2024.*?(?:TOTAL|Total).*?\$?([\d,]+\.?\d*)'
        total_match = re.search(total_pattern, full_text)
        if total_match and not total_amount:
            total_amount = total_match.group(1).replace(',', '')
        
        # Look for paid date pattern
        date_pattern = r'2024.*?PAID\s+(\d{1,2}/\d{1,2}/\d{4})'
        date_match = re.search(date_pattern, full_text)
        if date_match and not date_paid:
            date_paid = date_match.group(1)
    
    return total_amount, date_paid

def process_tax_document(file_path: str) -> Tuple[bool, str, Optional[str], Optional[str]]:
    """
    Process a tax document (PDF or DOCX) and extract 2024 tax information.
    
    Returns:
        Tuple of (success, message, total_amount, date_paid)
    """
    import os
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return extract_tax_info_from_pdf(file_path)
    elif file_ext == '.docx':
        try:
            doc = Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    text += '\n' + row_text
            
            if not text.strip():
                return False, "No text extracted from document", None, None
            
            total_amount, date_paid = parse_tax_text(text)
            
            if total_amount or date_paid:
                return True, "Successfully extracted tax information", total_amount, date_paid
            else:
                return False, "Could not find 2024 tax information in document", None, None
                
        except Exception as e:
            return False, f"Error processing Word document: {str(e)}", None, None
    else:
        return False, f"Unsupported file type: {file_ext}", None, None