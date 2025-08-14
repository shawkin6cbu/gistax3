import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinterdnd2 import DND_FILES
from desoto.services.title_chain import (
    process_title_document
)
from desoto.services.tax_document import process_tax_document
import threading
import os
from docx import Document
import re

class ProcessingTab(ttk.Frame):
    def __init__(self, parent, shared_data):
        super().__init__(parent, padding=20)
        self.shared_data = shared_data

        # Main layout grid
        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=1)

        # --- LEFT COLUMN ---
        left_frame = ttk.Frame(self)
        left_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        left_frame.columnconfigure(0, weight=1)

        # --- Property Information ---
        prop_frame = ttk.LabelFrame(left_frame, text="Property Information", padding=15)
        prop_frame.grid(row=0, column=0, sticky="ew", pady=(0, 20))
        prop_frame.columnconfigure(1, weight=1)
        
        self.pin_var = tk.StringVar()
        self.address_var = tk.StringVar()
        self.owner_var = tk.StringVar()
        self.city_var = tk.StringVar()
        self.legal_desc_var = tk.StringVar()
        
        self._create_entry_row(prop_frame, "PIN:", self.pin_var, 0)
        self._create_entry_row(prop_frame, "Address:", self.address_var, 1)
        self._create_entry_row(prop_frame, "Owner:", self.owner_var, 2)
        self._create_entry_row(prop_frame, "City/State/ZIP:", self.city_var, 3)
        self._create_entry_row(prop_frame, "Legal Desc:", self.legal_desc_var, 4)

        # --- Tax Information ---
        tax_frame = ttk.LabelFrame(left_frame, text="Tax Information", padding=15)
        tax_frame.grid(row=1, column=0, sticky="ew", pady=(0, 20))
        tax_frame.columnconfigure(1, weight=1)
        
        self.tax_2024_total_var = tk.StringVar()
        self.tax_2024_paid_var = tk.StringVar(value="PAID")
        self.tax_2024_date_paid_var = tk.StringVar()
        self.tax_2025_est_var = tk.StringVar()
        
        self._create_entry_row(tax_frame, "2024 Total:", self.tax_2024_total_var, 0)
        self._create_entry_row(tax_frame, "Date Paid:", self.tax_2024_date_paid_var, 1)
        self._create_entry_row(tax_frame, "2024 Status:", self.tax_2024_paid_var, 2, is_combo=True)
        self._create_entry_row(tax_frame, "2025 Estimated:", self.tax_2025_est_var, 3)

        # --- Tax Document Drop Zone ---
        tax_doc_frame = ttk.LabelFrame(left_frame, text="Tax Document", padding=15)
        tax_doc_frame.grid(row=2, column=0, sticky="ew", pady=(0, 20))
        tax_doc_frame.columnconfigure(1, weight=1)
        
        self.tax_doc_var = tk.StringVar()
        self.tax_doc_entry = self._create_file_drop_row(tax_doc_frame, "Tax Doc:", self.tax_doc_var, self.browse_tax_document, self._drop_on_tax_doc)

        # --- RIGHT COLUMN ---
        right_frame = ttk.Frame(self)
        right_frame.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        right_frame.columnconfigure(0, weight=1)

        # --- Title Chain Document ---
        doc_frame = ttk.LabelFrame(right_frame, text="Title Chain Document", padding=15)
        doc_frame.grid(row=0, column=0, sticky="ew", pady=(0, 20))
        doc_frame.columnconfigure(1, weight=1)
        
        self.title_doc_var = tk.StringVar()
        self.title_doc_entry = self._create_file_drop_row(doc_frame, "File:", self.title_doc_var, self.browse_title_document, self._drop_on_doc)
        
        # --- Title Chain Summary ---
        title_frame = ttk.LabelFrame(right_frame, text="Title Chain Summary", padding=15)
        title_frame.grid(row=1, column=0, sticky="ew", pady=(0, 20))
        title_frame.columnconfigure(0, weight=1)
        
        self.title_summary_var = tk.StringVar(value="Not loaded yet.")
        ttk.Label(title_frame, textvariable=self.title_summary_var, wraplength=400, justify="left").grid(row=0, column=0, columnspan=2, sticky="w")
        ttk.Button(title_frame, text="View & Edit Details", command=self.view_title_details).grid(row=1, column=0, sticky="w", pady=(10,0))
        
        # --- Document Details ---
        doc_details_frame = ttk.LabelFrame(right_frame, text="Document Details", padding=15)
        doc_details_frame.grid(row=2, column=0, sticky="ew", pady=(0, 20))
        doc_details_frame.columnconfigure(1, weight=1)
        self.lender_var = tk.StringVar()
        self.borrower_var = tk.StringVar()
        self._create_entry_row(doc_details_frame, "Lender:", self.lender_var, 0)
        self._create_entry_row(doc_details_frame, "Borrower:", self.borrower_var, 1)

        # --- Document Generation ---
        doc_gen_frame = ttk.LabelFrame(right_frame, text="Document Generation", padding=15)
        doc_gen_frame.grid(row=3, column=0, sticky="ew")
        doc_gen_frame.columnconfigure(1, weight=1)
        
        self.output_path_var = tk.StringVar()
        self._create_entry_row(doc_gen_frame, "Output Path:", self.output_path_var, 0, browse_btn=True)
        
        self.generate_btn = ttk.Button(doc_gen_frame, text="Generate Document", command=self.generate_document)
        self.generate_btn.grid(row=1, column=1, sticky="w", pady=(10,0))
        
        self.progress = ttk.Progressbar(doc_gen_frame, mode="indeterminate")
        self.progress.grid(row=2, column=0, columnspan=3, sticky="ew", pady=(10,0))

    def _create_entry_row(self, parent, label_text, var, row, browse_btn=False, is_combo=False):
        ttk.Label(parent, text=label_text).grid(row=row, column=0, sticky="e", padx=(0, 10), pady=5)
        
        if is_combo:
            combo = ttk.Combobox(
                parent,
                textvariable=var,
                values=["PAID", "UNPAID", "PARTIAL"],
                state="normal",
                width=28
            )
            combo.grid(row=row, column=1, sticky="ew", pady=5)
        else:
            entry = ttk.Entry(parent, textvariable=var, width=40)
            entry.grid(row=row, column=1, sticky="ew", pady=5, padx=(0, 5 if browse_btn else 0))
        
        if browse_btn:
            ttk.Button(parent, text="...", width=4, command=self.browse_output).grid(row=row, column=2, pady=5)
        
        parent.columnconfigure(1, weight=1)

    def _create_file_drop_row(self, parent, label_text, var, browse_cmd, drop_cmd):
        parent.drop_target_register(DND_FILES)
        parent.dnd_bind("<<Drop>>", drop_cmd)

        ttk.Label(parent, text=label_text).grid(row=0, column=0, sticky="w", pady=5)
        
        entry = ttk.Entry(parent, textvariable=var, width=30)
        entry.grid(row=0, column=1, sticky="ew", pady=5, padx=(10, 5))
        
        ttk.Button(parent, text="Browse", command=browse_cmd).grid(row=0, column=2, sticky="e", pady=5)

        # Auto-process when file path changes
        var.trace_add('write', lambda *_: self.process_document_if_valid(var, drop_cmd))
        return entry

    def process_document_if_valid(self, var, process_func):
        path = (var.get() or "").strip()
        if path and os.path.exists(path) and path.lower().endswith(('.pdf', '.docx')):
            process_func()

    def _drop_on_tax_doc(self, event=None):
        if event:
            path = event.data.strip('{}')
            if path.lower().endswith(('.pdf', '.docx')):
                self.tax_doc_var.set(path)
        self.process_tax_document()

    def _drop_on_doc(self, event=None):
        if event:
            path = event.data.strip('{}')
            if path.lower().endswith(('.pdf', '.docx')):
                self.title_doc_var.set(path)
        self.process_title_document()

    def browse_tax_document(self):
        path = filedialog.askopenfilename(
            title="Select Tax Document", 
            filetypes=[("PDF files","*.pdf"), ("Word documents", "*.docx"), ("All files","*.*")]
        )
        if path:
            self.tax_doc_var.set(path)

    def process_tax_document(self):
        file_path = self.tax_doc_var.get().strip()
        if not file_path or not os.path.exists(file_path):
            return
        
        self.tax_2024_total_var.set("Processing...")
        self.tax_2024_date_paid_var.set("Processing...")
        
        threading.Thread(target=self._process_tax_document_thread, args=(file_path,), daemon=True).start()

    def _process_tax_document_thread(self, file_path):
        try:
            success, msg, total_amount, date_paid = process_tax_document(file_path)
            
            def update_ui():
                if success:
                    self.tax_2024_total_var.set(total_amount or "")
                    self.tax_2024_date_paid_var.set(date_paid or "")
                    if date_paid:
                        self.tax_2024_paid_var.set("PAID")
                    self.shared_data.update_data({
                        "tax_2024_total": total_amount or "",
                        "tax_2024_date_paid": date_paid or "",
                        "tax_2024_paid_status": "PAID" if date_paid else self.tax_2024_paid_var.get()
                    })
                    messagebox.showinfo("Success", f"Tax information extracted:\n2024 Total: ${total_amount}\nDate Paid: {date_paid}")
                else:
                    self.tax_2024_total_var.set("")
                    self.tax_2024_date_paid_var.set("")
                    messagebox.showwarning("Processing Error", msg)
            self.after(0, update_ui)
            
        except Exception as e:
            error_msg = f"Error processing tax document: {str(e)}"
            self.after(0, lambda: messagebox.showerror("Processing Error", error_msg))

    def browse_title_document(self):
        path = filedialog.askopenfilename(
            title="Select Title Chain Document", 
            filetypes=[("PDF files","*.pdf"), ("Word documents", "*.docx"), ("All files","*.*")]
        )
        if path:
            self.title_doc_var.set(path)

    def process_title_document(self):
        file_path = self.title_doc_var.get().strip()
        if not file_path or not os.path.exists(file_path):
            return
        
        self.title_summary_var.set("Processing document...")
        threading.Thread(target=self._process_title_document_thread, args=(file_path,), daemon=True).start()

    def _process_title_document_thread(self, file_path):
        try:
            success, msg, chain_deeds, all_entries = process_title_document(file_path)
            
            def update_ui():
                if success:
                    self.shared_data.set_data("title_chain_kept", chain_deeds)
                    self.shared_data.set_data("title_chain_all", all_entries)
                    summary = f"Found {len(chain_deeds)} vesting deeds in 24-month chain."
                    self.title_summary_var.set(summary)
                else:
                    self.shared_data.set_data("title_chain_kept", [])
                    self.shared_data.set_data("title_chain_all", [])
                    self.title_summary_var.set(msg)
                    messagebox.showerror("Processing Error", msg)
            self.after(0, update_ui)
            
        except Exception as e:
            error_msg = f"Error processing document: {str(e)}"
            self.after(0, lambda: self.title_summary_var.set("Processing failed."))
            self.after(0, lambda: messagebox.showerror("Processing Error", error_msg))

    def load_from_tabs(self):
        self.pin_var.set(self.shared_data.get_data("parcel_pin") or "")
        self.address_var.set(self.shared_data.get_data("parcel_address") or "")
        self.owner_var.set(self.shared_data.get_data("parcel_owner") or "")
        self.city_var.set(self.shared_data.get_data("parcel_city_state_zip") or "")
        self.legal_desc_var.set(self.shared_data.get_data("parcel_legal_description") or "")
        
        self.tax_2024_total_var.set(self.shared_data.get_data("tax_2024_total") or "")
        self.tax_2024_paid_var.set(self.shared_data.get_data("tax_2024_paid_status") or "PAID")
        self.tax_2024_date_paid_var.set(self.shared_data.get_data("tax_2024_date_paid") or "")
        self.tax_2025_est_var.set(self.shared_data.get_data("tax_2025_estimated") or "")

        results = self.shared_data.get_data("title_chain_kept")
        self.title_summary_var.set(f"{len(results)} vesting deeds found." if results else "No title chain data found.")

    def view_title_details(self):
        all_entries = self.shared_data.get_data("title_chain_all") or []
        kept_entries = self.shared_data.get_data("title_chain_kept") or []

        if not all_entries:
            messagebox.showinfo("Title Chain Details", "No title chain data has been loaded.")
            return

        details_win = tk.Toplevel(self)
        details_win.title("Manage Title Chain Entries")
        details_win.geometry("900x700")

        paned_window = ttk.PanedWindow(details_win, orient=tk.VERTICAL)
        paned_window.pack(fill="both", expand=True, padx=10, pady=10)

        cols = ("Date", "Grantor", "Grantee", "Instrument", "Book-Page")
        sort_states = {'keep': {'col': 'Date', 'rev': True}, 'other': {'col': 'Date', 'rev': True}}

        def sort_tree(tree, tree_key, col):
            state = sort_states[tree_key]
            reverse = not state['rev'] if state['col'] == col else False
            state.update({'col': col, 'rev': reverse})

            items = [(tree.set(child, col), child) for child in tree.get_children('')]
            
            if col == 'Date':
                from datetime import datetime
                items.sort(key=lambda it: datetime.strptime(it[0], "%m/%d/%Y"), reverse=reverse)
            elif col == 'Book-Page':
                items.sort(key=lambda it: [int(p) for p in re.match(r'(\d+)-(\d+)', it[0]).groups()] if re.match(r'(\d+)-(\d+)', it[0]) else [0,0], reverse=reverse)
            else:
                items.sort(key=lambda it: it[0].lower(), reverse=reverse)

            for index, (_, child) in enumerate(items):
                tree.move(child, '', index)

            for c in cols:
                arrow = ' ↓' if reverse else ' ↑'
                tree.heading(c, text=c + (arrow if c == col else ''))

        def create_tree_view(parent, key):
            frame = ttk.LabelFrame(parent, text="Kept" if key == 'keep' else "Other Entries", padding=10)
            parent.add(frame, weight=1 if key == 'keep' else 2)
            tree = ttk.Treeview(frame, columns=cols, show="headings")
            for col in cols:
                tree.heading(col, text=col, command=lambda c=col, t=tree, k=key: sort_tree(t, k, c))
                tree.column(col, width=150, anchor="w")
            tree.pack(fill="both", expand=True)
            return tree

        keep_tree = create_tree_view(paned_window, 'keep')
        other_tree = create_tree_view(paned_window, 'other')

        all_entries_map = {(e.date_string, e.book_page, e.instrument): e for e in all_entries}
        
        def populate_trees():
            kept_set = {(e.date_string, e.book_page, e.instrument) for e in kept_entries}
            
            for tree in [keep_tree, other_tree]:
                for item in tree.get_children():
                    tree.delete(item)

            for entry in sorted(all_entries, key=lambda e: e.date, reverse=True):
                values = (entry.date_string, entry.grantor, entry.grantee, entry.instrument, entry.book_page)
                key = (entry.date_string, entry.book_page, entry.instrument)
                (keep_tree if key in kept_set else other_tree).insert("", "end", values=values)

            sort_tree(keep_tree, 'keep', 'Date')
            sort_tree(other_tree, 'other', 'Date')

        def move_item(from_tree, to_tree, event):
            selected_item = from_tree.selection()
            if not selected_item: return
            item_values = from_tree.item(selected_item, "values")
            to_tree.insert("", "end", values=item_values)
            from_tree.delete(selected_item)

        keep_tree.bind("<Double-1>", lambda e: move_item(keep_tree, other_tree, e))
        other_tree.bind("<Double-1>", lambda e: move_item(other_tree, keep_tree, e))

        def on_close():
            new_kept_keys = {tuple(keep_tree.item(child, "values")) for child in keep_tree.get_children()}
            new_kept_entries = [all_entries_map.get((v[0], v[4], v[3])) for v in new_kept_keys]
            new_kept_entries = [e for e in new_kept_entries if e] # Filter out None
            new_kept_entries.sort(key=lambda e: e.date, reverse=True)
            
            self.shared_data.set_data("title_chain_kept", new_kept_entries)
            self.title_summary_var.set(f"{len(new_kept_entries)} vesting deeds in chain.")
            details_win.destroy()

        details_win.protocol("WM_DELETE_WINDOW", on_close)
        populate_trees()

    def browse_output(self):
        path = filedialog.asksaveasfilename(
            title="Save Document As",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
        )
        if path:
            self.output_path_var.set(path)

    def generate_document(self):
        self.sync_to_shared_data()
        doc_path = (self.title_doc_var.get() or "").strip()
        out_path = self.output_path_var.get().strip()

        output_path = out_path or os.path.abspath("TitleDocs.docx")
        if not out_path and doc_path:
            try:
                output_path = os.path.join(os.path.dirname(doc_path), "TitleDocs.docx")
            except Exception:
                pass

        self.generate_btn.config(state="disabled")
        self.progress.start()
        threading.Thread(target=self._generate_document_thread, args=(output_path,), daemon=True).start()

    def _generate_document_thread(self, output_path):
        try:
            success, msg = self._create_full_document(output_path)
            self.after(0, lambda: messagebox.showinfo("Success", msg) if success else messagebox.showerror("Error", msg))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", f"An unexpected error has occurred: {e}"))
        finally:
            self.after(0, self.progress.stop)
            self.after(0, lambda: self.generate_btn.config(state="normal"))

    def get_template_path(self):
        try:
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            path = os.path.join(base_dir, "templates", "td_tmplt2.docx")
            return path if os.path.exists(path) else None
        except Exception:
            return None

    def sync_to_shared_data(self):
        self.shared_data.update_data({
            "parcel_pin": self.pin_var.get(),
            "parcel_address": self.address_var.get(),
            "parcel_owner": self.owner_var.get(),
            "parcel_city_state_zip": self.city_var.get(),
            "parcel_legal_description": self.legal_desc_var.get(),
            "tax_2024_total": self.tax_2024_total_var.get(),
            "tax_2024_paid_status": self.tax_2024_paid_var.get(),
            "tax_2024_date_paid": self.tax_2024_date_paid_var.get(),
            "tax_2025_estimated": self.tax_2025_est_var.get(),
            "lender": self.lender_var.get(),
            "borrower": self.borrower_var.get(),
        })

    def _create_full_document(self, output_path):
        template_path = self.get_template_path()
        if not template_path:
            return False, "Template file 'td_tmplt2.docx' not found."

        doc_path = (self.title_doc_var.get() or "").strip()
        if doc_path and os.path.exists(doc_path) and not self.shared_data.get_data("title_chain_kept"):
            success, _, chain_deeds, all_entries = process_title_document(doc_path)
            if success:
                self.shared_data.set_data("title_chain_kept", chain_deeds)
                self.shared_data.set_data("title_chain_all", all_entries)

        doc = Document(template_path)

        def smart_title_case(text):
            if not isinstance(text, str) or not text or text.strip().isdigit() or '$' in text:
                return text
            
            preserve_upper = {'LLC', 'PLLC', 'INC', 'CO', 'CORP', 'LP', 'LLP', 'PA', 'PC', 'LTD', 'II', 'III', 'IV', 'JR', 'SR', 'MS', 'US'}
            return ' '.join([word.upper() if word.strip('.,;:').upper() in preserve_upper else word.capitalize() for word in text.split()])

        values_map = {
            "{{{PARCEL}}}": self.pin_var.get(),
            "{{{PROPSTRE}}}": smart_title_case(self.address_var.get()),
            "{{{SLRLAST}}}": smart_title_case(self.owner_var.get()),
            "{{{CITY_STATE_ZIP}}}": self.city_var.get(),
            "{{{LEGAL_DESC}}}": self.legal_desc_var.get(),
            "{{{TAXAMT}}}": f"${self.tax_2024_total_var.get()}" if self.tax_2024_total_var.get() else "",
            "{{{TAXDAT}}}": self.tax_2024_date_paid_var.get(),
            "{{{TAX_2025_EST}}}": f"${self.tax_2025_est_var.get()}" if self.tax_2025_est_var.get() else "",
            "{{{Lender}}}": smart_title_case(self.lender_var.get()),
            "{{{BYRLAST}}}": smart_title_case(self.borrower_var.get()),
            "{{{LOAN_AMOUNT}}}": "",
        }

        for p in doc.paragraphs:
            for run in p.runs:
                for key, value in values_map.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)

        for table in doc.tables:
            is_chain_table = 'GRANTOR' in ' '.join(cell.text.upper() for cell in table.rows[0].cells)
            if not is_chain_table:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                for key, value in values_map.items():
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)

        chain_deeds = self.shared_data.get_data("title_chain_kept") or []
        chain_table = next((t for t in doc.tables if 'GRANTOR' in ' '.join(c.text.upper() for c in t.rows[0].cells)), None)

        if chain_table and chain_deeds:
            while len(chain_table.rows) > 1:
                chain_table._element.remove(chain_table.rows[-1]._element)
            for deed in chain_deeds:
                cells = chain_table.add_row().cells
                cells[0].text, cells[1].text, cells[2].text, cells[3].text, cells[4].text = deed.grantor, deed.grantee, deed.instrument, deed.date_string, deed.book_page

        doc.save(output_path)
        return True, f"Document successfully generated at:\n{output_path}"
